"""
Test script to verify all 9 elements of a 3x3 matrix are visible in Word.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def make_omml_matrix(rows, bracket='['):
    beg_chr = '|' if bracket == '|' else '['
    end_chr = '|' if bracket == '|' else ']'
    mr_xml_list = []
    for row in rows:
        e_xml_list = []
        for cell in row:
            e_xml_list.append(f'<m:e><m:r><m:rPr><m:scr m:val="roman"/><m:sty m:val="p"/></m:rPr><m:t>{cell}</m:t></m:r></m:e>')
        mr_xml_list.append(f'<m:mr>{"".join(e_xml_list)}</m:mr>')
        
    m_body = "".join(mr_xml_list)
    xml = f"""
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:d>
        <m:dPr>
          <m:begChr m:val="{beg_chr}"/>
          <m:endChr m:val="{end_chr}"/>
          <m:grow m:val="on"/>
        </m:dPr>
        <m:e>
          <m:m>
            <m:mPr>
              <m:baseJc m:val="center"/>
              <m:plcHide m:val="on"/>
            </m:mPr>
            {m_body}
          </m:m>
        </m:e>
      </m:d>
    </m:oMath>
    """
    return parse_xml(xml)

doc = Document()
# Set Normal style with SINGLE spacing (not Pt(14) fixed)
style = doc.styles['Normal']
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(2)
pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
pf.line_spacing = None

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p.paragraph_format.line_spacing = None
p.paragraph_format.space_after = Pt(4)
r = p.add_run("6. The value of x for which the matrix A = ")
p._element.append(make_omml_matrix([['0', '1', '-2'], ['-1', '0', 'x'], ['2', '-3', '0']], bracket='['))
p.add_run(" is skew-symmetric, is:\n(a) 3   (b) -3   (c) 2   (d) -2")

p2 = doc.add_paragraph()
p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
p2.paragraph_format.line_spacing = None
p2.paragraph_format.space_after = Pt(4)
p2.add_run("28. Using minors and cofactors, evaluate the determinant:\n    |A| = ")
p2._element.append(make_omml_matrix([['1', '2', '3'], ['0', '1', '4'], ['5', '6', '0']], bracket='|'))

doc.save(r"C:\Users\rajan\.gemini\antigravity\scratch\test_matrix_fixed.docx")
print("Saved test_matrix_fixed.docx successfully!")
