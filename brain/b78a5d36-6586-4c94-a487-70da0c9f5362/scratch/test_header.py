"""
Test official DAV header in python-docx.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders: top, bottom, left, right.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>\n'
                          f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'  <w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>\n'
                      f'  <w:top w:w="{top}" w:type="dxa"/>\n'
                      f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
                      f'  <w:left w:w="{left}" w:type="dxa"/>\n'
                      f'  <w:right w:w="{right}" w:type="dxa"/>\n'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

doc = Document()
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

cell = table.cell(0, 0)
cell.width = Cm(18.4)
set_cell_border(cell)
set_cell_margins(cell, top=160, bottom=160, left=240, right=240)

# Paragraph 1: Title
p1 = cell.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(2)
p1.paragraph_format.space_after = Pt(2)
p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r1 = p1.add_run("DAV INSTITUTIONS, JHARKHAND ZONE-F")
r1.bold = True
r1.font.name = 'Times New Roman'
r1.font.size = Pt(13.5)

# Paragraph 2: Subtitle
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(6)
p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r2 = p2.add_run("Half Yearly Examination, 2026-27")
r2.font.name = 'Times New Roman'
r2.font.size = Pt(11.5)

# Paragraph 3: Class & Full Marks
p3 = cell.add_paragraph()
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(2)
p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r3_l = p3.add_run("Class: XII")
r3_l.bold = True
r3_l.font.name = 'Times New Roman'
r3_l.font.size = Pt(11)

# Tab stop or spaces
p3.add_run("\t\t\t\t\t\t\t\t\t")
r3_r = p3.add_run("FULL MARKS: 80")
r3_r.bold = True
r3_r.font.name = 'Times New Roman'
r3_r.font.size = Pt(11)

# Paragraph 4: Subject & Time Allowed
p4 = cell.add_paragraph()
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after = Pt(2)
p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
r4_l = p4.add_run("Subject: Mathematics (Set A)")
r4_l.bold = True
r4_l.font.name = 'Times New Roman'
r4_l.font.size = Pt(11)

p4.add_run("\t\t\t\t\t\t\t")
r4_r = p4.add_run("TIME ALLOWED: 3 HOURS")
r4_r.bold = True
r4_r.font.name = 'Times New Roman'
r4_r.font.size = Pt(11)

# Centered Asterisks below box
p_ast = doc.add_paragraph()
p_ast.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ast.paragraph_format.space_before = Pt(6)
p_ast.paragraph_format.space_after = Pt(6)
r_ast = p_ast.add_run("************")
r_ast.bold = True
r_ast.font.name = 'Times New Roman'
r_ast.font.size = Pt(12)

doc.save(r"C:\Users\rajan\.gemini\antigravity\scratch\test_header.docx")
print("Header test generated successfully!")
