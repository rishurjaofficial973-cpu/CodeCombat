"""
Generate complete Set A and Set B Question Papers and Marking Schemes
with embedded high-contrast images and proper picture paragraph formatting (preventing clipping).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT

OUT_DIR = r"C:\Users\rajan\.gemini\antigravity\scratch\math_exam_2026"
IMG_DIR = os.path.join(OUT_DIR, "images")


def create_base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(14)
    return doc


def add_p(doc, text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    if left_indent > 0:
        p.paragraph_format.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_image(doc, img_name, width_inches=3.6):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # CRITICAL: Reset line_spacing to SINGLE so Word does not clip the image to 14pt!
        p.paragraph_format.line_spacing = None
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))


def build_set_a_qp():
    doc = create_base_doc()
    
    # Header
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "CLASS XII — MATHEMATICS (SET A)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("Time Allowed: 3 Hours")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    p.add_run("\t\t\t\t\t\t\t\t")
    r2 = p.add_run("Maximum Marks: 80")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    # Instructions
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        add_p(doc, f"{i}. {inst}", size=11, space_after=1, left_indent=0.5)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # SECTION A
    add_p(doc, "SECTION A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 1 to 18 carry 1 mark each. Questions 19 and 20 are Assertion–Reason based, 1 mark each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_a = [
        ("1. Let A = {1, 2, 3} and R = {(1,1), (2,2), (3,3), (1,2)} be a relation on A. Then R is:\n(a) An equivalence relation\t\t\t(b) Reflexive and transitive only\n(c) Reflexive and symmetric only\t\t(d) Symmetric and transitive only"),
        ("2. The function f : ℝ → ℝ defined by f(x) = x² is:\n(a) One-one and onto\t\t\t(b) One-one but not onto\n(c) Onto but not one-one\t\t\t(d) Neither one-one nor onto"),
        ("3. The principal value of sin⁻¹(−1/2) is:\n(a) π/6\t\t(b) −π/6\t\t(c) 5π/6\t\t(d) −5π/6"),
        ("4. The domain of cos⁻¹(2x − 1) is:\n(a) [0, 1]\t\t(b) [−1, 1]\t\t(c) [−1, 0]\t\t(d) [0, 2]"),
        ("5. For any square matrix A, the matrix (A + Aᵀ) is always:\n(a) Skew-symmetric\t\t\t(b) Symmetric\n(c) A diagonal matrix\t\t\t(d) A null matrix"),
        ("6. The value of x for which the matrix A = [[0, 1, -2], [-1, 0, x], [2, -3, 0]] is skew-symmetric, is:\n(a) 3\t\t(b) −3\t\t(c) 2\t\t(d) −2"),
        ("7. If A is any square matrix, then the matrix (A − Aᵀ) is always:\n(a) Symmetric\t(b) Skew-symmetric\t(c) Null matrix\t(d) Identity matrix"),
        ("8. If A is a 3 × 3 matrix such that |A| = 5, then the value of |3A| is:\n(a) 15\t\t(b) 45\t\t(c) 135\t\t(d) 405"),
        ("9. If A is a square matrix of order 3 with |A| = 4, then |adj A| equals:\n(a) 4\t\t(b) 8\t\t(c) 16\t\t(d) 64"),
        ("10. The function f(x) = |x| is:\n(a) Continuous and differentiable everywhere\n(b) Continuous everywhere but not differentiable at x = 0\n(c) Discontinuous at x = 0\n(d) Differentiable everywhere but not continuous"),
        ("11. If y = e²ˣ, then dy/dx equals:\n(a) e²ˣ\t\t(b) 2e²ˣ\t\t(c) 2x e²ˣ\t\t(d) eˣ"),
        ("12. If f(x) = kx + 1 for x ≤ 5 and f(x) = 3x − 5 for x > 5 is continuous at x = 5, then k =:\n(a) 9/5\t\t(b) 5/9\t\t(c) 2\t\t(d) 3"),
        ("13. The rate of change of the area of a circle with respect to its radius r, at r = 5 cm, is:\n(a) 5π cm²/cm\t(b) 10π cm²/cm\t(c) 25π cm²/cm\t(d) 2π cm²/cm"),
        ("14. The function f(x) = x² − 4x + 6 is strictly increasing in the interval:\n(a) (−∞, 2)\t\t(b) (2, ∞)\t\t(c) (−2, 2)\t\t(d) ℝ"),
        ("15. The maximum value of the function f(x) = sin x + cos x on the interval [0, π/2] is:\n(a) 1\t\t(b) √2\t\t(c) 2\t\t(d) 1/√2"),
        ("16. The total revenue (in ₹) from the sale of x units of a product is R(x) = 3x² + 36x + 5. The marginal revenue when x = 15 is:\n(a) ₹ 116\t\t(b) ₹ 96\t\t(c) ₹ 90\t\t(d) ₹ 126"),
        ("17. ∫ sec²x dx equals:\n(a) tan x + C\t(b) −cot x + C\t(c) sec x tan x + C\t(d) cot x + C"),
        ("18. ∫ 1/(1 + x²) dx equals:\n(a) tan⁻¹x + C\t(b) sin⁻¹x + C\t(c) cot⁻¹x + C\t(d) sec⁻¹x + C")
    ]
    for q in q_sec_a:
        add_p(doc, q, space_after=3)
    
    # Q19 AR
    add_p(doc, "19. The graph of a function y = f(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_a_q19.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = |x − 2| shown in the graph is continuous at x = 2.\nReason (R): The function f(x) = |x − 2| is differentiable at x = 2.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=3)

    # Q20 AR
    add_p(doc, "20. The graph of a function y = f(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_a_q20.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The point (3, 4) marked on the graph of f(x) = −(x − 3)² + 4 is a point of local maximum.\nReason (R): At a point of local maximum, f′(x) = 0 and f″(x) < 0.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=5)

    # SECTION B
    add_p(doc, "SECTION B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 21 to 25 carry 2 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_b = [
        "21. Check whether the relation R = {(1,1), (2,2), (3,3), (1,2), (2,3)} on the set A = {1, 2, 3} is reflexive, symmetric and transitive. Justify your answer.",
        "22. If A = [[1, 2], [3, 4]] and B = [[2, 0], [1, 3]], find 2A − 3B.",
        "23. Differentiate y = sin(x² + 1) with respect to x.",
        "24. An edge of a variable cube is increasing at the rate of 3 cm/s. How fast is the volume of the cube increasing when the edge is 10 cm long?",
        "25. Evaluate: ∫ x e^(x²) dx."
    ]
    for q in q_sec_b:
        add_p(doc, q, space_after=3)

    # SECTION C
    add_p(doc, "SECTION C", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 26 to 31 carry 3 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_c = [
        "26. Find the value of:\n(i) sin⁻¹(sin(2π/3))\t\t(ii) cos⁻¹(cos(7π/6))\t\t(iii) tan⁻¹(tan(3π/4))",
        "27. Find the value of: tan⁻¹(1) + cos⁻¹(−1/2) + sin⁻¹(−1/2).",
        "28. Using minors and cofactors, evaluate the determinant:\n    | 1   2   3 |\n    | 0   1   4 |\n    | 5   6   0 |",
        "29. If y = xˣ (x > 0), find dy/dx.",
        "30. Find the intervals in which the function f(x) = 2x³ − 15x² + 36x + 1 is (i) strictly increasing (ii) strictly decreasing.",
        "31. Evaluate ∫ x sin x dx using integration by parts."
    ]
    for q in q_sec_c:
        add_p(doc, q, space_after=3)

    # SECTION D
    add_p(doc, "SECTION D", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 32 to 35 carry 5 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_d = [
        "32. Express the matrix A = [[1, 2, -3], [3, -1, 2], [-2, 1, 4]] as the sum of a symmetric and a skew-symmetric matrix.",
        "33. Using the matrix method, solve the following system of linear equations:\n    x + y + z = 6\n    y + 3z = 11\n    x − 2y + z = 0",
        "34. If y = (sin⁻¹x)², prove that (1 − x²) d²y/dx² − x dy/dx − 2 = 0.",
        "35. Evaluate: ∫ (x² + 1)/(x² − 5x + 6) dx."
    ]
    for q in q_sec_d:
        add_p(doc, q, space_after=3)

    # SECTION E
    add_p(doc, "SECTION E", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 36 to 38 are case-based questions of 4 marks each, with sub-parts of 2, 1 and 1 marks.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q36 Case Study 1
    add_p(doc, "36. Case Study 1 — Relations and Functions", bold=True, size=12, space_after=2)
    add_p(doc, "Two friends, Rohan and Simran, are studying equivalence relations on sets. Consider the set A = {1, 2, 3, 4, 5} and the relation R on A defined by R = {(a, b) : |a − b| is even}. The diagram below illustrates the partition of set A into equivalence classes under relation R:", space_after=2)
    add_image(doc, "set_a_q36.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Show that R is an equivalence relation on A.\t\t(2 marks)\n(ii) Write the equivalence class of the element 2.\t\t(1 mark)\n(iii) Is (1, 5) ∈ R? Justify your answer.\t\t(1 mark)", space_after=4)

    # Q37 Case Study 2
    add_p(doc, "37. Case Study 2 — Application of Derivatives", bold=True, size=12, space_after=2)
    add_p(doc, "A manufacturing firm analyzes its production costs. The total cost of producing x units of a product is modelled by the cost function C(x) = 0.005x³ − 0.02x² + 30x + 5000 (in ₹). The total cost curve is depicted below:", space_after=2)
    add_image(doc, "set_a_q37.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the marginal cost function, MC(x) = dC/dx.\t\t(2 marks)\n(ii) Find the marginal cost when x = 3 units.\t\t(1 mark)\n(iii) State, with reason, whether the cost is increasing or decreasing at x = 3.\t\t(1 mark)", space_after=4)

    # Q38 Case Study 3
    add_p(doc, "38. Case Study 3 — Integrals", bold=True, size=12, space_after=2)
    add_p(doc, "A particle moves along a straight line such that its velocity at time t seconds is given by v(t) = 3t² − 12t + 9 (in m/s). It is given that the particle starts at the origin, i.e. s(0) = 0. The velocity–time graph of the motion is shown below:", space_after=2)
    add_image(doc, "set_a_q38.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the displacement function s(t).\t\t(2 marks)\n(ii) Find the displacement of the particle at t = 2 s.\t\t(1 mark)\n(iii) Find the value(s) of t ∈ [0, 3] at which the particle is momentarily at rest.\t\t(1 mark)", space_after=5)

    add_p(doc, "— End of Question Paper —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_docx = os.path.join(OUT_DIR, "Set_A_Question_Paper.docx")
    doc.save(out_docx)
    print(f"Generated {out_docx}")


def build_set_b_qp():
    doc = create_base_doc()
    
    # Header
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "CLASS XII — MATHEMATICS (SET B)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("Time Allowed: 3 Hours")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    p.add_run("\t\t\t\t\t\t\t\t")
    r2 = p.add_run("Maximum Marks: 80")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    # Instructions
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        add_p(doc, f"{i}. {inst}", size=11, space_after=1, left_indent=0.5)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # SECTION A
    add_p(doc, "SECTION A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 1 to 18 carry 1 mark each. Questions 19 and 20 are Assertion–Reason based, 1 mark each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_a = [
        ("1. Let R = {(1,1), (2,2), (3,3), (1,3), (3,1)} be a relation on A = {1, 2, 3}. Then R is:\n(a) Reflexive only\t\t\t(b) Symmetric only\n(c) An equivalence relation\t\t(d) Reflexive and symmetric only"),
        ("2. The number of all one-one functions from the set {1, 2, 3} to itself is:\n(a) 3\t\t(b) 6\t\t(c) 9\t\t(d) 27"),
        ("3. The principal value of cos⁻¹(−1/√2) is:\n(a) π/4\t\t(b) 3π/4\t\t(c) −π/4\t\t(d) 5π/4"),
        ("4. The range of the principal value branch of tan⁻¹x is:\n(a) [−π/2, π/2]\t\t(b) (−π/2, π/2)\t\t(c) [0, π]\t\t(d) (0, π)"),
        ("5. If A and B are square matrices of the same order, then (A + B)ᵀ equals:\n(a) Aᵀ + Bᵀ\t\t(b) Aᵀ − Bᵀ\t\t(c) AᵀBᵀ\t\t(d) BᵀAᵀ"),
        ("6. If the matrix A = [[0, a, -3], [2, 0, -1], [b, 1, 0]] is skew-symmetric, then the value of (a + b) is:\n(a) 1\t\t(b) −1\t\t(c) 5\t\t(d) −5"),
        ("7. A matrix A = [aᵢⱼ]₃ₓ₃ is a diagonal matrix if:\n(a) aᵢⱼ = 0 for all i, j\t\t\t(b) aᵢⱼ = 0 for i ≠ j\n(c) aᵢⱼ = 0 for i = j\t\t\t(d) aᵢⱼ ≠ 0 for all i, j"),
        ("8. If A is a 3 × 3 matrix with |A| = −2, then |3A| equals:\n(a) −6\t\t(b) −54\t\t(c) 54\t\t(d) −18"),
        ("9. The area of the triangle with vertices (1, 0), (6, 0) and (4, 3), found using determinants, is:\n(a) 15 sq units\t\t(b) 15/2 sq units\t\t(c) 7 sq units\t\t(d) 12 sq units"),
        ("10. If y = sin(3x), then dy/dx is:\n(a) cos(3x)\t\t(b) 3cos(3x)\t\t(c) −3cos(3x)\t\t(d) 3sin(3x)"),
        ("11. The function f(x) = x³ is:\n(a) Continuous but not differentiable at x = 0\n(b) Differentiable but not continuous at x = 0\n(c) Both continuous and differentiable everywhere on ℝ\n(d) Neither continuous nor differentiable at x = 0"),
        ("12. The value of k for which f(x) = kx² (for x ≤ 2) and f(x) = 3 (for x > 2) is continuous at x = 2, is:\n(a) 3/4\t\t(b) 4/3\t\t(c) 3\t\t(d) 4"),
        ("13. The rate of change of the volume V = (4/3)πr³ of a sphere with respect to its radius r is:\n(a) 4πr²\t\t(b) 2πr\t\t(c) (4/3)πr²\t\t(d) 4πr"),
        ("14. The function f(x) = x³ − 6x² + 9x + 15 is strictly decreasing in the interval:\n(a) (−∞, 1)\t\t(b) (1, 3)\t\t(c) (3, ∞)\t\t(d) (1, ∞)"),
        ("15. The rate of change of the total surface area S = 4πr² of a sphere with respect to its radius r, when r = 3 cm, is:\n(a) 12π cm²/cm\t(b) 24π cm²/cm\t(c) 36π cm²/cm\t(d) 6π cm²/cm"),
        ("16. The function f(x) = 2x³ − 3x² − 12x + 4 has a local maximum at:\n(a) x = 2\t\t(b) x = −1\t\t(c) x = 0\t\t(d) x = 1"),
        ("17. ∫ cos x dx equals:\n(a) sin x + C\t(b) −sin x + C\t(c) cos x + C\t(d) −cos x + C"),
        ("18. ∫ 1/√(1 − x²) dx equals:\n(a) sin⁻¹x + C\t(b) cos⁻¹x + C\t(c) tan⁻¹x + C\t(d) sec⁻¹x + C")
    ]
    for q in q_sec_a:
        add_p(doc, q, space_after=3)

    # Q19 AR
    add_p(doc, "19. The graph of the greatest integer function y = [x] is shown below:", bold=False, space_after=2)
    add_image(doc, "set_b_q19.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = [x] (greatest integer function) is discontinuous at every integer.\nReason (R): For any integer n, lim(x→n) [x] does not exist.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=3)

    # Q20 AR
    add_p(doc, "20. The graph of the logarithmic function y = ln(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_b_q20.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = log x is strictly increasing on (0, ∞).\nReason (R): f′(x) = 1/x > 0 for all x ∈ (0, ∞).\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=5)

    # SECTION B
    add_p(doc, "SECTION B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 21 to 25 carry 2 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_b = [
        "21. Let A = {1, 2, 3} and R = {(1,1), (2,2), (3,3), (1,3), (3,1)} be a relation on A. Is R an equivalence relation? Justify your answer.",
        "22. If A = [[3, 1], [-1, 2]] and B = [[1, 0], [-1, 2]], find A + 2B.",
        "23. Differentiate y = cos(x²) with respect to x.",
        "24. The radius of a circle is increasing uniformly at the rate of 3 cm/s. Find the rate at which the area of the circle is increasing when the radius is 10 cm.",
        "25. Evaluate: ∫ e^(sin x) · cos x dx."
    ]
    for q in q_sec_b:
        add_p(doc, q, space_after=3)

    # SECTION C
    add_p(doc, "SECTION C", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 26 to 31 carry 3 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_c = [
        "26. Find the value of:\n(i) cos⁻¹(cos(5π/4))\t\t(ii) sin⁻¹(sin(−π/3))\t\t(iii) tan⁻¹(tan(5π/6))",
        "27. Find the domain of:\n(i) f(x) = sin⁻¹(3x − 1)\t\t(ii) g(x) = cos⁻¹((x − 1)/2)\nAlso, find the principal value of tan⁻¹(−1).",
        "28. Evaluate the determinant:\n    |  2  -1   3 |\n    |  1   2  -1 |\n    |  0   3   2 |",
        "29. If y = x^(sin x) (x > 0), find dy/dx.",
        "30. Find the intervals in which the function f(x) = x³ − 6x² + 9x + 15 is (i) strictly increasing (ii) strictly decreasing.",
        "31. Evaluate ∫ x cos x dx using integration by parts."
    ]
    for q in q_sec_c:
        add_p(doc, q, space_after=3)

    # SECTION D
    add_p(doc, "SECTION D", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 32 to 35 carry 5 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    q_sec_d = [
        "32. Express the matrix A = [[3, -2, 5], [4, 1, -3], [-1, 2, 1]] as the sum of a symmetric and a skew-symmetric matrix.",
        "33. Using the matrix method, solve the following system of linear equations:\n    x − y + 2z = 7\n    3x + 4y − 5z = −5\n    2x − y + 3z = 12",
        "34. If y = (tan⁻¹x)², show that (1 + x²)² d²y/dx² + 2x(1 + x²) dy/dx = 2.",
        "35. Evaluate: ∫ (x² + x + 1)/((x + 1)(x + 2)) dx."
    ]
    for q in q_sec_d:
        add_p(doc, q, space_after=3)

    # SECTION E
    add_p(doc, "SECTION E", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 36 to 38 are case-based questions of 4 marks each, with sub-parts of 2, 1 and 1 marks.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q36 Case Study 1
    add_p(doc, "36. Case Study 1 — Relations and Functions", bold=True, size=12, space_after=2)
    add_p(doc, "A mathematics club defines a relation R on the set A = {1, 2, 3, 4, 6} as R = {(a, b) : a divides b, where a, b ∈ A}. The directed graph below illustrates the divisibility relation among the elements:", space_after=2)
    add_image(doc, "set_b_q36.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Show that R is reflexive and transitive.\t\t(2 marks)\n(ii) Is R symmetric? Justify with an example.\t\t(1 mark)\n(iii) Is R an equivalence relation? Give reason.\t\t(1 mark)", space_after=4)

    # Q37 Case Study 2
    add_p(doc, "37. Case Study 2 — Application of Derivatives", bold=True, size=12, space_after=2)
    add_p(doc, "A manufacturer's total revenue (in ₹) from the sale of x units of a commodity is given by R(x) = 13x² + 26x + 15. The revenue function curve is shown below:", space_after=2)
    add_image(doc, "set_b_q37.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the marginal revenue function, MR(x) = dR/dx.\t\t(2 marks)\n(ii) Find the marginal revenue when x = 7 units.\t\t(1 mark)\n(iii) State, with reason, whether the revenue is increasing or decreasing at x = 7.\t\t(1 mark)", space_after=4)

    # Q38 Case Study 3
    add_p(doc, "38. Case Study 3 — Integrals", bold=True, size=12, space_after=2)
    add_p(doc, "A car starts from rest and its velocity at time t seconds is given by v(t) = 6t² − 4t (in m/s). It is given that the displacement at t = 0 is zero, i.e. s(0) = 0. The velocity–time graph of the motion is shown below:", space_after=2)
    add_image(doc, "set_b_q38.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the displacement function s(t).\t\t(2 marks)\n(ii) Find the displacement at t = 3 s.\t\t(1 mark)\n(iii) Find the value of t (t > 0) at which the car is momentarily at rest.\t\t(1 mark)", space_after=5)

    add_p(doc, "— End of Question Paper —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_docx = os.path.join(OUT_DIR, "Set_B_Question_Paper.docx")
    doc.save(out_docx)
    print(f"Generated {out_docx}")


def build_marking_schemes():
    # Build Set A MS Docx
    doc_a = create_base_doc()
    add_p(doc_a, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc_a, "CLASS XII — MATHEMATICS (SET A)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc_a, "MARKING SCHEME & STEP-BY-STEP SOLUTIONS", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc_a, "Time Allowed: 3 Hours\t\t\t\t\t\t\t\tMaximum Marks: 80", bold=True, size=12, space_after=3)
    add_p(doc_a, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    with open(os.path.join(OUT_DIR, "Set_A_Marking_Scheme.txt"), 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if line.startswith("SECTION "):
                add_p(doc_a, line, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            elif line.startswith("Q") and ("Case Study" in line or "Solution:" in line or line.endswith(".")):
                add_p(doc_a, line, bold=True, size=12, space_after=2)
            elif line.startswith("Correct Option:"):
                add_p(doc_a, line, bold=True, size=12, space_after=3)
            elif line.startswith("Final Answer:"):
                add_p(doc_a, line, bold=True, size=12, space_after=3)
            elif line.startswith("— End"):
                add_p(doc_a, line, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            else:
                add_p(doc_a, line, size=11, space_after=1)
    
    out_a = os.path.join(OUT_DIR, "Set_A_Marking_Scheme.docx")
    doc_a.save(out_a)
    print(f"Generated {out_a}")

    # Build Set B MS Docx
    doc_b = create_base_doc()
    add_p(doc_b, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc_b, "CLASS XII — MATHEMATICS (SET B)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc_b, "MARKING SCHEME & STEP-BY-STEP SOLUTIONS", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc_b, "Time Allowed: 3 Hours\t\t\t\t\t\t\t\tMaximum Marks: 80", bold=True, size=12, space_after=3)
    add_p(doc_b, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    with open(os.path.join(OUT_DIR, "Set_B_Marking_Scheme.txt"), 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if line.startswith("SECTION "):
                add_p(doc_b, line, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            elif line.startswith("Q") and ("Case Study" in line or "Solution:" in line or line.endswith(".")):
                add_p(doc_b, line, bold=True, size=12, space_after=2)
            elif line.startswith("Correct Option:"):
                add_p(doc_b, line, bold=True, size=12, space_after=3)
            elif line.startswith("Final Answer:"):
                add_p(doc_b, line, bold=True, size=12, space_after=3)
            elif line.startswith("— End"):
                add_p(doc_b, line, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
            else:
                add_p(doc_b, line, size=11, space_after=1)
                
    out_b = os.path.join(OUT_DIR, "Set_B_Marking_Scheme.docx")
    doc_b.save(out_b)
    print(f"Generated {out_b}")


if __name__ == '__main__':
    build_set_a_qp()
    build_set_b_qp()
    build_marking_schemes()
    print("All documents generated successfully.")
