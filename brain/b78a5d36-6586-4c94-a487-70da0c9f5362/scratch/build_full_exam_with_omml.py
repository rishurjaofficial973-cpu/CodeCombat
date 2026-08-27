"""
Generate complete Set A and Set B Question Papers and Marking Schemes
with:
1. Native MS Word OMML equation brackets for all matrices (3x3, 2x2, 3x1 vectors) and determinants.
2. Single Line Spacing (line_spacing = None, line_spacing_rule = SINGLE) so Word dynamically expands line height and displays all 3 rows (all 9 elements) of every matrix and determinant with full brackets.
3. Centered, high-contrast 300 DPI diagrams with solid white backgrounds and no clipping.
4. 100% NCERT rationalized syllabus compliance.
5. Strict Times New Roman 12pt, A4 portrait, narrow margin formatting with no box separators.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml

OUT_DIR = r"C:\Users\rajan\.gemini\antigravity\scratch\math_exam_2026"
IMG_DIR = os.path.join(OUT_DIR, "images")


def make_omml_matrix(rows, bracket='['):
    """
    rows: 2D list of strings, e.g. [['1', '2'], ['3', '4']]
    bracket: '[' for square bracket matrix, '|' for determinant, '(' for paren
    """
    beg_chr = '|' if bracket == '|' else ('(' if bracket == '(' else '[')
    end_chr = '|' if bracket == '|' else (')' if bracket == '(' else ']')
    
    mr_xml_list = []
    for row in rows:
        e_xml_list = []
        for cell in row:
            e_xml_list.append(f'<m:e><m:r><m:rPr><m:scr m:val="roman"/><m:sty m:val="p"/></m:rPr><m:t>{cell}</m:t></m:r></m:e>')
        mr_xml_list.append(f'<m:mr>{"".join(e_xml_list)}</m:mr>')
        
    m_body = "".join(mr_xml_list)
    
    xml = f"""
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:d>
        <m:dPr>
          <m:begChr m:val="{beg_chr}"/>
          <m:endChr m:val="{end_chr}"/>
          <m:grow m:val="on"/>
        </m:dPr>
        <m:e>
          <m:m>
            <m:mPr>
              <m:baseJc m:val="center"/>
              <m:plcHide m:val="on"/>
            </m:mPr>
            {m_body}
          </m:m>
        </m:e>
      </m:d>
    </m:oMath>
    """
    return parse_xml(xml)


def create_base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    # Use SINGLE spacing so Word dynamically expands for multi-line math objects (3x3 matrices)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    return doc


def add_p(doc, text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = None
    if left_indent > 0:
        p.paragraph_format.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_matrix_p(doc, text_before="", rows=None, bracket='[', text_after="", space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = None
    if text_before:
        r1 = p.add_run(text_before)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0, 0, 0)
    if rows:
        p._element.append(make_omml_matrix(rows, bracket=bracket))
    if text_after:
        r2 = p.add_run(text_after)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_image(doc, img_name, width_inches=3.6):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = None
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))


def build_set_a_qp():
    doc = create_base_doc()
    
    # Header
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "CLASS XII — MATHEMATICS (SET A)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("Time Allowed: 3 Hours")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    p.add_run("\t\t\t\t\t\t\t\t")
    r2 = p.add_run("Maximum Marks: 80")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    # Instructions
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        add_p(doc, f"{i}. {inst}", size=11, space_after=1, left_indent=0.5)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # SECTION A
    add_p(doc, "SECTION A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 1 to 18 carry 1 mark each. Questions 19 and 20 are Assertion–Reason based, 1 mark each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "1. Let A = {1, 2, 3} and R = {(1,1), (2,2), (3,3), (1,2)} be a relation on A. Then R is:\n(a) An equivalence relation\t\t\t(b) Reflexive and transitive only\n(c) Reflexive and symmetric only\t\t(d) Symmetric and transitive only", space_after=3)
    add_p(doc, "2. The function f : ℝ → ℝ defined by f(x) = x² is:\n(a) One-one and onto\t\t\t(b) One-one but not onto\n(c) Onto but not one-one\t\t\t(d) Neither one-one nor onto", space_after=3)
    add_p(doc, "3. The principal value of sin⁻¹(−1/2) is:\n(a) π/6\t\t(b) −π/6\t\t(c) 5π/6\t\t(d) −5π/6", space_after=3)
    add_p(doc, "4. The domain of cos⁻¹(2x − 1) is:\n(a) [0, 1]\t\t(b) [−1, 1]\t\t(c) [−1, 0]\t\t(d) [0, 2]", space_after=3)
    add_p(doc, "5. For any square matrix A, the matrix (A + Aᵀ) is always:\n(a) Skew-symmetric\t\t\t(b) Symmetric\n(c) A diagonal matrix\t\t\t(d) A null matrix", space_after=3)
    
    # Q6 with OMML 3x3 Matrix
    add_matrix_p(doc, "6. The value of x for which the matrix A = ",
                 [['0', '1', '-2'], ['-1', '0', 'x'], ['2', '-3', '0']], '[',
                 " is skew-symmetric, is:\n(a) 3\t\t(b) −3\t\t(c) 2\t\t(d) −2", space_after=3)

    add_p(doc, "7. If A is any square matrix, then the matrix (A − Aᵀ) is always:\n(a) Symmetric\t(b) Skew-symmetric\t(c) Null matrix\t(d) Identity matrix", space_after=3)
    add_p(doc, "8. If A is a 3 × 3 matrix such that |A| = 5, then the value of |3A| is:\n(a) 15\t\t(b) 45\t\t(c) 135\t\t(d) 405", space_after=3)
    add_p(doc, "9. If A is a square matrix of order 3 with |A| = 4, then |adj A| equals:\n(a) 4\t\t(b) 8\t\t(c) 16\t\t(d) 64", space_after=3)
    add_p(doc, "10. The function f(x) = |x| is:\n(a) Continuous and differentiable everywhere\n(b) Continuous everywhere but not differentiable at x = 0\n(c) Discontinuous at x = 0\n(d) Differentiable everywhere but not continuous", space_after=3)
    add_p(doc, "11. If y = e²ˣ, then dy/dx equals:\n(a) e²ˣ\t\t(b) 2e²ˣ\t\t(c) 2x e²ˣ\t\t(d) eˣ", space_after=3)
    add_p(doc, "12. If f(x) = kx + 1 for x ≤ 5 and f(x) = 3x − 5 for x > 5 is continuous at x = 5, then k =:\n(a) 9/5\t\t(b) 5/9\t\t(c) 2\t\t(d) 3", space_after=3)
    add_p(doc, "13. The rate of change of the area of a circle with respect to its radius r, at r = 5 cm, is:\n(a) 5π cm²/cm\t(b) 10π cm²/cm\t(c) 25π cm²/cm\t(d) 2π cm²/cm", space_after=3)
    add_p(doc, "14. The function f(x) = x² − 4x + 6 is strictly increasing in the interval:\n(a) (−∞, 2)\t\t(b) (2, ∞)\t\t(c) (−2, 2)\t\t(d) ℝ", space_after=3)
    add_p(doc, "15. The maximum value of the function f(x) = sin x + cos x on the interval [0, π/2] is:\n(a) 1\t\t(b) √2\t\t(c) 2\t\t(d) 1/√2", space_after=3)
    add_p(doc, "16. The total revenue (in ₹) from the sale of x units of a product is R(x) = 3x² + 36x + 5. The marginal revenue when x = 15 is:\n(a) ₹ 116\t\t(b) ₹ 96\t\t(c) ₹ 90\t\t(d) ₹ 126", space_after=3)
    add_p(doc, "17. ∫ sec²x dx equals:\n(a) tan x + C\t(b) −cot x + C\t(c) sec x tan x + C\t(d) cot x + C", space_after=3)
    add_p(doc, "18. ∫ 1/(1 + x²) dx equals:\n(a) tan⁻¹x + C\t(b) sin⁻¹x + C\t(c) cot⁻¹x + C\t(d) sec⁻¹x + C", space_after=3)

    # Q19 AR
    add_p(doc, "19. The graph of a function y = f(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_a_q19.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = |x − 2| shown in the graph is continuous at x = 2.\nReason (R): The function f(x) = |x − 2| is differentiable at x = 2.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=3)

    # Q20 AR
    add_p(doc, "20. The graph of a function y = f(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_a_q20.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The point (3, 4) marked on the graph of f(x) = −(x − 3)² + 4 is a point of local maximum.\nReason (R): At a point of local maximum, f′(x) = 0 and f″(x) < 0.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=5)

    # SECTION B
    add_p(doc, "SECTION B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 21 to 25 carry 2 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "21. Check whether the relation R = {(1,1), (2,2), (3,3), (1,2), (2,3)} on the set A = {1, 2, 3} is reflexive, symmetric and transitive. Justify your answer.", space_after=3)
    
    # Q22 with OMML 2x2 Matrices
    p22 = doc.add_paragraph()
    p22.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p22.paragraph_format.line_spacing = None
    p22.paragraph_format.space_after = Pt(3)
    r = p22.add_run("22. If A = ")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p22._element.append(make_omml_matrix([['1', '2'], ['3', '4']], bracket='['))
    r2 = p22.add_run(" and B = ")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    p22._element.append(make_omml_matrix([['2', '0'], ['1', '3']], bracket='['))
    r3 = p22.add_run(", find 2A − 3B.")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

    add_p(doc, "23. Differentiate y = sin(x² + 1) with respect to x.", space_after=3)
    add_p(doc, "24. An edge of a variable cube is increasing at the rate of 3 cm/s. How fast is the volume of the cube increasing when the edge is 10 cm long?", space_after=3)
    add_p(doc, "25. Evaluate: ∫ x e^(x²) dx.", space_after=3)

    # SECTION C
    add_p(doc, "SECTION C", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 26 to 31 carry 3 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "26. Find the value of:\n(i) sin⁻¹(sin(2π/3))\t\t(ii) cos⁻¹(cos(7π/6))\t\t(iii) tan⁻¹(tan(3π/4))", space_after=3)
    add_p(doc, "27. Find the value of: tan⁻¹(1) + cos⁻¹(−1/2) + sin⁻¹(−1/2).", space_after=3)
    
    # Q28 with OMML 3x3 Determinant Bracket
    add_matrix_p(doc, "28. Using minors and cofactors, evaluate the determinant:\n    |A| = ",
                 [['1', '2', '3'], ['0', '1', '4'], ['5', '6', '0']], '|', "", space_after=3)

    add_p(doc, "29. If y = xˣ (x > 0), find dy/dx.", space_after=3)
    add_p(doc, "30. Find the intervals in which the function f(x) = 2x³ − 15x² + 36x + 1 is (i) strictly increasing (ii) strictly decreasing.", space_after=3)
    add_p(doc, "31. Evaluate ∫ x sin x dx using integration by parts.", space_after=3)

    # SECTION D
    add_p(doc, "SECTION D", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 32 to 35 carry 5 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q32 with OMML 3x3 Matrix
    add_matrix_p(doc, "32. Express the matrix A = ",
                 [['1', '2', '-3'], ['3', '-1', '2'], ['-2', '1', '4']], '[',
                 " as the sum of a symmetric and a skew-symmetric matrix.", space_after=3)

    add_p(doc, "33. Using the matrix method, solve the following system of linear equations:\n    x + y + z = 6\n    y + 3z = 11\n    x − 2y + z = 0", space_after=3)
    add_p(doc, "34. If y = (sin⁻¹x)², prove that (1 − x²) d²y/dx² − x dy/dx − 2 = 0.", space_after=3)
    add_p(doc, "35. Evaluate: ∫ (x² + 1)/(x² − 5x + 6) dx.", space_after=3)

    # SECTION E
    add_p(doc, "SECTION E", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 36 to 38 are case-based questions of 4 marks each, with sub-parts of 2, 1 and 1 marks.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q36 Case Study 1
    add_p(doc, "36. Case Study 1 — Relations and Functions", bold=True, size=12, space_after=2)
    add_p(doc, "Two friends, Rohan and Simran, are studying equivalence relations on sets. Consider the set A = {1, 2, 3, 4, 5} and the relation R on A defined by R = {(a, b) : |a − b| is even}. The diagram below illustrates the partition of set A into equivalence classes under relation R:", space_after=2)
    add_image(doc, "set_a_q36.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Show that R is an equivalence relation on A.\t\t(2 marks)\n(ii) Write the equivalence class of the element 2.\t\t(1 mark)\n(iii) Is (1, 5) ∈ R? Justify your answer.\t\t(1 mark)", space_after=4)

    # Q37 Case Study 2
    add_p(doc, "37. Case Study 2 — Application of Derivatives", bold=True, size=12, space_after=2)
    add_p(doc, "A manufacturing firm analyzes its production costs. The total cost of producing x units of a product is modelled by the cost function C(x) = 0.005x³ − 0.02x² + 30x + 5000 (in ₹). The total cost curve is depicted below:", space_after=2)
    add_image(doc, "set_a_q37.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the marginal cost function, MC(x) = dC/dx.\t\t(2 marks)\n(ii) Find the marginal cost when x = 3 units.\t\t(1 mark)\n(iii) State, with reason, whether the cost is increasing or decreasing at x = 3.\t\t(1 mark)", space_after=4)

    # Q38 Case Study 3
    add_p(doc, "38. Case Study 3 — Integrals", bold=True, size=12, space_after=2)
    add_p(doc, "A particle moves along a straight line such that its velocity at time t seconds is given by v(t) = 3t² − 12t + 9 (in m/s). It is given that the particle starts at the origin, i.e. s(0) = 0. The velocity–time graph of the motion is shown below:", space_after=2)
    add_image(doc, "set_a_q38.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the displacement function s(t).\t\t(2 marks)\n(ii) Find the displacement of the particle at t = 2 s.\t\t(1 mark)\n(iii) Find the value(s) of t ∈ [0, 3] at which the particle is momentarily at rest.\t\t(1 mark)", space_after=5)

    add_p(doc, "— End of Question Paper —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_docx = os.path.join(OUT_DIR, "Set_A_Question_Paper.docx")
    doc.save(out_docx)
    print(f"Generated {out_docx}")


def build_set_b_qp():
    doc = create_base_doc()
    
    # Header
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "CLASS XII — MATHEMATICS (SET B)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run("Time Allowed: 3 Hours")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    p.add_run("\t\t\t\t\t\t\t\t")
    r2 = p.add_run("Maximum Marks: 80")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    # Instructions
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        add_p(doc, f"{i}. {inst}", size=11, space_after=1, left_indent=0.5)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # SECTION A
    add_p(doc, "SECTION A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 1 to 18 carry 1 mark each. Questions 19 and 20 are Assertion–Reason based, 1 mark each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "1. Let R = {(1,1), (2,2), (3,3), (1,3), (3,1)} be a relation on A = {1, 2, 3}. Then R is:\n(a) Reflexive only\t\t\t(b) Symmetric only\n(c) An equivalence relation\t\t(d) Reflexive and symmetric only", space_after=3)
    add_p(doc, "2. The number of all one-one functions from the set {1, 2, 3} to itself is:\n(a) 3\t\t(b) 6\t\t(c) 9\t\t(d) 27", space_after=3)
    add_p(doc, "3. The principal value of cos⁻¹(−1/√2) is:\n(a) π/4\t\t(b) 3π/4\t\t(c) −π/4\t\t(d) 5π/4", space_after=3)
    add_p(doc, "4. The range of the principal value branch of tan⁻¹x is:\n(a) [−π/2, π/2]\t\t(b) (−π/2, π/2)\t\t(c) [0, π]\t\t(d) (0, π)", space_after=3)
    add_p(doc, "5. If A and B are square matrices of the same order, then (A + B)ᵀ equals:\n(a) Aᵀ + Bᵀ\t\t(b) Aᵀ − Bᵀ\t\t(c) AᵀBᵀ\t\t(d) BᵀAᵀ", space_after=3)
    
    # Q6 with OMML 3x3 Matrix
    add_matrix_p(doc, "6. If the matrix A = ",
                 [['0', 'a', '-3'], ['2', '0', '-1'], ['b', '1', '0']], '[',
                 " is skew-symmetric, then the value of (a + b) is:\n(a) 1\t\t(b) −1\t\t(c) 5\t\t(d) −5", space_after=3)

    add_p(doc, "7. A matrix A = [aᵢⱼ]₃ₓ₃ is a diagonal matrix if:\n(a) aᵢⱼ = 0 for all i, j\t\t\t(b) aᵢⱼ = 0 for i ≠ j\n(c) aᵢⱼ = 0 for i = j\t\t\t(d) aᵢⱼ ≠ 0 for all i, j", space_after=3)
    add_p(doc, "8. If A is a 3 × 3 matrix with |A| = −2, then |3A| equals:\n(a) −6\t\t(b) −54\t\t(c) 54\t\t(d) −18", space_after=3)
    add_p(doc, "9. The area of the triangle with vertices (1, 0), (6, 0) and (4, 3), found using determinants, is:\n(a) 15 sq units\t\t(b) 15/2 sq units\t\t(c) 7 sq units\t\t(d) 12 sq units", space_after=3)
    add_p(doc, "10. If y = sin(3x), then dy/dx is:\n(a) cos(3x)\t\t(b) 3cos(3x)\t\t(c) −3cos(3x)\t\t(d) 3sin(3x)", space_after=3)
    add_p(doc, "11. The function f(x) = x³ is:\n(a) Continuous but not differentiable at x = 0\n(b) Differentiable but not continuous at x = 0\n(c) Both continuous and differentiable everywhere on ℝ\n(d) Neither continuous nor differentiable at x = 0", space_after=3)
    add_p(doc, "12. The value of k for which f(x) = kx² (for x ≤ 2) and f(x) = 3 (for x > 2) is continuous at x = 2, is:\n(a) 3/4\t\t(b) 4/3\t\t(c) 3\t\t(d) 4", space_after=3)
    add_p(doc, "13. The rate of change of the volume V = (4/3)πr³ of a sphere with respect to its radius r is:\n(a) 4πr²\t\t(b) 2πr\t\t(c) (4/3)πr²\t\t(d) 4πr", space_after=3)
    add_p(doc, "14. The function f(x) = x³ − 6x² + 9x + 15 is strictly decreasing in the interval:\n(a) (−∞, 1)\t\t(b) (1, 3)\t\t(c) (3, ∞)\t\t(d) (1, ∞)", space_after=3)
    add_p(doc, "15. The rate of change of the total surface area S = 4πr² of a sphere with respect to its radius r, when r = 3 cm, is:\n(a) 12π cm²/cm\t(b) 24π cm²/cm\t(c) 36π cm²/cm\t(d) 6π cm²/cm", space_after=3)
    add_p(doc, "16. The function f(x) = 2x³ − 3x² − 12x + 4 has a local maximum at:\n(a) x = 2\t\t(b) x = −1\t\t(c) x = 0\t\t(d) x = 1", space_after=3)
    add_p(doc, "17. ∫ cos x dx equals:\n(a) sin x + C\t(b) −sin x + C\t(c) cos x + C\t(d) −cos x + C", space_after=3)
    add_p(doc, "18. ∫ 1/√(1 − x²) dx equals:\n(a) sin⁻¹x + C\t(b) cos⁻¹x + C\t(c) tan⁻¹x + C\t(d) sec⁻¹x + C", space_after=3)

    # Q19 AR
    add_p(doc, "19. The graph of the greatest integer function y = [x] is shown below:", bold=False, space_after=2)
    add_image(doc, "set_b_q19.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = [x] (greatest integer function) is discontinuous at every integer.\nReason (R): For any integer n, lim(x→n) [x] does not exist.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=3)

    # Q20 AR
    add_p(doc, "20. The graph of the logarithmic function y = ln(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_b_q20.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = log x is strictly increasing on (0, ∞).\nReason (R): f′(x) = 1/x > 0 for all x ∈ (0, ∞).\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=5)

    # SECTION B
    add_p(doc, "SECTION B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 21 to 25 carry 2 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "21. Let A = {1, 2, 3} and R = {(1,1), (2,2), (3,3), (1,3), (3,1)} be a relation on A. Is R an equivalence relation? Justify your answer.", space_after=3)
    
    # Q22 with OMML 2x2 Matrices
    p22 = doc.add_paragraph()
    p22.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p22.paragraph_format.line_spacing = None
    p22.paragraph_format.space_after = Pt(3)
    r = p22.add_run("22. If A = ")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p22._element.append(make_omml_matrix([['3', '1'], ['-1', '2']], bracket='['))
    r2 = p22.add_run(" and B = ")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    p22._element.append(make_omml_matrix([['1', '0'], ['-1', '2']], bracket='['))
    r3 = p22.add_run(", find A + 2B.")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

    add_p(doc, "23. Differentiate y = cos(x²) with respect to x.", space_after=3)
    add_p(doc, "24. The radius of a circle is increasing uniformly at the rate of 3 cm/s. Find the rate at which the area of the circle is increasing when the radius is 10 cm.", space_after=3)
    add_p(doc, "25. Evaluate: ∫ e^(sin x) · cos x dx.", space_after=3)

    # SECTION C
    add_p(doc, "SECTION C", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 26 to 31 carry 3 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "26. Find the value of:\n(i) cos⁻¹(cos(5π/4))\t\t(ii) sin⁻¹(sin(−π/3))\t\t(iii) tan⁻¹(tan(5π/6))", space_after=3)
    add_p(doc, "27. Find the domain of:\n(i) f(x) = sin⁻¹(3x − 1)\t\t(ii) g(x) = cos⁻¹((x − 1)/2)\nAlso, find the principal value of tan⁻¹(−1).", space_after=3)
    
    # Q28 with OMML 3x3 Determinant
    add_matrix_p(doc, "28. Evaluate the determinant:\n    |A| = ",
                 [['2', '-1', '3'], ['1', '2', '-1'], ['0', '3', '2']], '|', "", space_after=3)

    add_p(doc, "29. If y = x^(sin x) (x > 0), find dy/dx.", space_after=3)
    add_p(doc, "30. Find the intervals in which the function f(x) = x³ − 6x² + 9x + 15 is (i) strictly increasing (ii) strictly decreasing.", space_after=3)
    add_p(doc, "31. Evaluate ∫ x cos x dx using integration by parts.", space_after=3)

    # SECTION D
    add_p(doc, "SECTION D", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 32 to 35 carry 5 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q32 with OMML 3x3 Matrix
    add_matrix_p(doc, "32. Express the matrix A = ",
                 [['3', '-2', '5'], ['4', '1', '-3'], ['-1', '2', '1']], '[',
                 " as the sum of a symmetric and a skew-symmetric matrix.", space_after=3)

    add_p(doc, "33. Using the matrix method, solve the following system of linear equations:\n    x − y + 2z = 7\n    3x + 4y − 5z = −5\n    2x − y + 3z = 12", space_after=3)
    add_p(doc, "34. If y = (tan⁻¹x)², show that (1 + x²)² d²y/dx² + 2x(1 + x²) dy/dx = 2.", space_after=3)
    add_p(doc, "35. Evaluate: ∫ (x² + x + 1)/((x + 1)(x + 2)) dx.", space_after=3)

    # SECTION E
    add_p(doc, "SECTION E", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 36 to 38 are case-based questions of 4 marks each, with sub-parts of 2, 1 and 1 marks.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q36 Case Study 1
    add_p(doc, "36. Case Study 1 — Relations and Functions", bold=True, size=12, space_after=2)
    add_p(doc, "A mathematics club defines a relation R on the set A = {1, 2, 3, 4, 6} as R = {(a, b) : a divides b, where a, b ∈ A}. The directed graph below illustrates the divisibility relation among the elements:", space_after=2)
    add_image(doc, "set_b_q36.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Show that R is reflexive and transitive.\t\t(2 marks)\n(ii) Is R symmetric? Justify with an example.\t\t(1 mark)\n(iii) Is R an equivalence relation? Give reason.\t\t(1 mark)", space_after=4)

    # Q37 Case Study 2
    add_p(doc, "37. Case Study 2 — Application of Derivatives", bold=True, size=12, space_after=2)
    add_p(doc, "A manufacturer's total revenue (in ₹) from the sale of x units of a commodity is given by R(x) = 13x² + 26x + 15. The revenue function curve is shown below:", space_after=2)
    add_image(doc, "set_b_q37.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the marginal revenue function, MR(x) = dR/dx.\t\t(2 marks)\n(ii) Find the marginal revenue when x = 7 units.\t\t(1 mark)\n(iii) State, with reason, whether the revenue is increasing or decreasing at x = 7.\t\t(1 mark)", space_after=4)

    # Q38 Case Study 3
    add_p(doc, "38. Case Study 3 — Integrals", bold=True, size=12, space_after=2)
    add_p(doc, "A car starts from rest and its velocity at time t seconds is given by v(t) = 6t² − 4t (in m/s). It is given that the displacement at t = 0 is zero, i.e. s(0) = 0. The velocity–time graph of the motion is shown below:", space_after=2)
    add_image(doc, "set_b_q38.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the displacement function s(t).\t\t(2 marks)\n(ii) Find the displacement at t = 3 s.\t\t(1 mark)\n(iii) Find the value of t (t > 0) at which the car is momentarily at rest.\t\t(1 mark)", space_after=5)

    add_p(doc, "— End of Question Paper —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_docx = os.path.join(OUT_DIR, "Set_B_Question_Paper.docx")
    doc.save(out_docx)
    print(f"Generated {out_docx}")


def build_set_a_ms():
    doc = create_base_doc()
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "CLASS XII — MATHEMATICS (SET A)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "MARKING SCHEME & STEP-BY-STEP SOLUTIONS", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Time Allowed: 3 Hours\t\t\t\t\t\t\t\tMaximum Marks: 80", bold=True, size=12, space_after=3)
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "SECTION A (20 Marks — 1 Mark Each)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    mcqs = [
        ("Q1. (b) Reflexive and transitive only\t\t[1 mark]\n    Reason: (1,1),(2,2),(3,3) ∈ R (reflexive). (1,2) ∈ R but (2,1) ∉ R (not symmetric). Transitive condition holds."),
        ("Q2. (d) Neither one-one nor onto\t\t[1 mark]\n    Reason: f(1)=f(-1)=1 (not 1-1). Range = [0, ∞) ≠ ℝ (codomain) (not onto)."),
        ("Q3. (b) −π/6\t\t[1 mark]\n    Reason: sin(-π/6) = -1/2 and -π/6 ∈ [-π/2, π/2]."),
        ("Q4. (a) [0, 1]\t\t[1 mark]\n    Reason: -1 ≤ 2x - 1 ≤ 1 ⇒ 0 ≤ 2x ≤ 2 ⇒ 0 ≤ x ≤ 1."),
        ("Q5. (b) Symmetric\t\t[1 mark]\n    Reason: (A + Aᵀ)ᵀ = Aᵀ + A = A + Aᵀ."),
        ("Q6. (a) 3\t\t[1 mark]\n    Reason: a₂₃ = -a₃₂ ⇒ x = -(-3) = 3."),
        ("Q7. (b) Skew-symmetric\t\t[1 mark]\n    Reason: (A - Aᵀ)ᵀ = Aᵀ - A = -(A - Aᵀ)."),
        ("Q8. (c) 135\t\t[1 mark]\n    Reason: |3A| = 3³ |A| = 27 × 5 = 135."),
        ("Q9. (c) 16\t\t[1 mark]\n    Reason: |adj A| = |A|ⁿ⁻¹ = 4³⁻¹ = 16."),
        ("Q10. (b) Continuous everywhere but not differentiable at x = 0\t\t[1 mark]\n    Reason: LHD at 0 is -1, RHD at 0 is +1."),
        ("Q11. (b) 2e²ˣ\t\t[1 mark]\n    Reason: d/dx(e²ˣ) = 2e²ˣ."),
        ("Q12. (a) 9/5\t\t[1 mark]\n    Reason: 5k + 1 = 3(5) - 5 = 10 ⇒ k = 9/5."),
        ("Q13. (b) 10π cm²/cm\t\t[1 mark]\n    Reason: dA/dr = 2πr = 10π at r = 5."),
        ("Q14. (b) (2, ∞)\t\t[1 mark]\n    Reason: f'(x) = 2x - 4 > 0 ⇒ x > 2."),
        ("Q15. (b) √2\t\t[1 mark]\n    Reason: f'(x) = cos x - sin x = 0 ⇒ x = π/4. f(π/4) = √2."),
        ("Q16. (d) ₹ 126\t\t[1 mark]\n    Reason: MR = dR/dx = 6x + 36. At x = 15: MR = 6(15) + 36 = 126."),
        ("Q17. (a) tan x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q18. (a) tan⁻¹x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q19. (c) A is true, but R is false\t\t[1 mark]\n    Reason: f(x) = |x - 2| is continuous everywhere, but not differentiable at the corner x = 2."),
        ("Q20. (a) Both A and R are true, and R is the correct explanation of A\t\t[1 mark]\n    Reason: f'(3) = 0 and f''(3) = -2 < 0 confirms local maximum at (3, 4).")
    ]
    for m in mcqs:
        add_p(doc, m, space_after=2)

    add_p(doc, "SECTION B (5 × 2 = 10 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    add_p(doc, "Q21. Check Reflexive, Symmetric, Transitive for R = {(1,1),(2,2),(3,3),(1,2),(2,3)}:\n- Reflexive: (1,1),(2,2),(3,3) ∈ R ⇒ Reflexive. [½ mark]\n- Symmetric: (1,2) ∈ R but (2,1) ∉ R ⇒ Not symmetric. [½ mark]\n- Transitive: (1,2) ∈ R and (2,3) ∈ R but (1,3) ∉ R ⇒ Not transitive. [1 mark]\nFinal Answer: R is reflexive, but neither symmetric nor transitive.", space_after=3)
    
    # Q22 MS with OMML 2x2
    p22 = doc.add_paragraph()
    p22.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p22.paragraph_format.line_spacing = None
    p22.paragraph_format.space_after = Pt(2)
    r = p22.add_run("Q22. 2A − 3B:\n2A = ")
    p22._element.append(make_omml_matrix([['2', '4'], ['6', '8']], bracket='['))
    r2 = p22.add_run(",  3B = ")
    p22._element.append(make_omml_matrix([['6', '0'], ['3', '9']], bracket='['))
    r3 = p22.add_run("  [1 mark]\n2A − 3B = ")
    p22._element.append(make_omml_matrix([['-4', '4'], ['3', '-1']], bracket='['))
    r4 = p22.add_run("  [1 mark]\nFinal Answer: ")
    p22._element.append(make_omml_matrix([['-4', '4'], ['3', '-1']], bracket='['))

    add_p(doc, "Q23. Differentiate y = sin(x² + 1):\ndy/dx = cos(x² + 1) · d/dx(x² + 1) [1 mark]\n      = 2x cos(x² + 1) [1 mark]\nFinal Answer: 2x cos(x² + 1)", space_after=3)
    add_p(doc, "Q24. Rate of increase of cube volume:\nV = x³ ⇒ dV/dt = 3x² (dx/dt) [1 mark]\nGiven dx/dt = 3 cm/s and x = 10 cm:\ndV/dt = 3(10)² (3) = 900 cm³/s [1 mark]\nFinal Answer: 900 cm³/s", space_after=3)
    add_p(doc, "Q25. Evaluate ∫ x e^(x²) dx:\nLet u = x² ⇒ du = 2x dx ⇒ x dx = du/2 [1 mark]\n∫ x e^(x²) dx = ½ ∫ eᵘ du = ½ eᵘ + C = ½ e^(x²) + C [1 mark]\nFinal Answer: ½ e^(x²) + C", space_after=3)

    add_p(doc, "SECTION C (6 × 3 = 18 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q26. Principal value evaluations:\n(i) sin⁻¹(sin(2π/3)) = sin⁻¹(sin(π − π/3)) = π/3 [1 mark]\n(ii) cos⁻¹(cos(7π/6)) = cos⁻¹(cos(2π − 5π/6)) = 5π/6 [1 mark]\n(iii) tan⁻¹(tan(3π/4)) = tan⁻¹(−tan(π/4)) = −π/4 [1 mark]", space_after=3)
    add_p(doc, "Q27. tan⁻¹(1) + cos⁻¹(−1/2) + sin⁻¹(−1/2):\ntan⁻¹(1) = π/4 [1 mark],  cos⁻¹(−1/2) = 2π/3 [½ mark],  sin⁻¹(−1/2) = −π/6 [½ mark]\nSum = π/4 + 2π/3 − π/6 = 9π/12 = 3π/4 [1 mark]\nFinal Answer: 3π/4", space_after=3)
    
    # Q28 MS with OMML 3x3 Determinant
    p28 = doc.add_paragraph()
    p28.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p28.paragraph_format.line_spacing = None
    p28.paragraph_format.space_after = Pt(2)
    r = p28.add_run("Q28. Determinant expansion along R₁ for |A| = ")
    p28._element.append(make_omml_matrix([['1', '2', '3'], ['0', '1', '4'], ['5', '6', '0']], bracket='|'))
    r2 = p28.add_run(":\n|A| = 1(0 − 24) − 2(0 − 20) + 3(0 − 5) [2 marks]\n    = −24 + 40 − 15 = 1 [1 mark]\nFinal Answer: 1")

    add_p(doc, "Q29. y = xˣ (x > 0):\nln y = x ln x [½ mark]\n(1/y) dy/dx = 1 + ln x [1.5 marks]\ndy/dx = xˣ(1 + ln x) [1 mark]\nFinal Answer: xˣ(1 + ln x)", space_after=3)
    add_p(doc, "Q30. Intervals of increase/decrease for f(x) = 2x³ − 15x² + 36x + 1:\nf'(x) = 6x² − 30x + 36 = 6(x − 2)(x − 3) [1 mark]\n- In (−∞, 2): f'(x) > 0 ⇒ Strictly Increasing [1 mark]\n- In (2, 3): f'(x) < 0 ⇒ Strictly Decreasing [½ mark]\n- In (3, ∞): f'(x) > 0 ⇒ Strictly Increasing [½ mark]\nFinal Answer: (i) Strictly increasing on (−∞, 2) ∪ (3, ∞); (ii) Strictly decreasing on (2, 3)", space_after=3)
    add_p(doc, "Q31. Evaluate ∫ x sin x dx using integration by parts:\nu = x ⇒ u' = 1;  v' = sin x ⇒ v = −cos x [1 mark]\n∫ x sin x dx = −x cos x − ∫ (−cos x) dx [1 mark]\n             = −x cos x + sin x + C [1 mark]\nFinal Answer: −x cos x + sin x + C", space_after=3)

    add_p(doc, "SECTION D (4 × 5 = 20 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    # Q32 MS with OMML 3x3 Matrices
    p32 = doc.add_paragraph()
    p32.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p32.paragraph_format.line_spacing = None
    p32.paragraph_format.space_after = Pt(2)
    r = p32.add_run("Q32. Express A = ")
    p32._element.append(make_omml_matrix([['1', '2', '-3'], ['3', '-1', '2'], ['-2', '1', '4']], bracket='['))
    r2 = p32.add_run(" as P + Q:\nAᵀ = ")
    p32._element.append(make_omml_matrix([['1', '3', '-2'], ['2', '-1', '1'], ['-3', '2', '4']], bracket='['))
    r3 = p32.add_run("  [1 mark]\nP = ½(A + Aᵀ) = ")
    p32._element.append(make_omml_matrix([['1', '5/2', '-5/2'], ['5/2', '-1', '3/2'], ['-5/2', '3/2', '4']], bracket='['))
    r4 = p32.add_run(" (Symmetric)  [1.5 marks]\nQ = ½(A − Aᵀ) = ")
    p32._element.append(make_omml_matrix([['0', '-1/2', '-1/2'], ['1/2', '0', '1/2'], ['1/2', '-1/2', '0']], bracket='['))
    r5 = p32.add_run(" (Skew-symmetric)  [1.5 marks]\nVerification: Pᵀ = P, Qᵀ = −Q, and P + Q = A.  [1 mark]")

    # Q33 MS with OMML System
    p33 = doc.add_paragraph()
    p33.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p33.paragraph_format.line_spacing = None
    p33.paragraph_format.space_after = Pt(2)
    r = p33.add_run("Q33. Solve system AX = B:\nA = ")
    p33._element.append(make_omml_matrix([['1', '1', '1'], ['0', '1', '3'], ['1', '-2', '1']], bracket='['))
    r2 = p33.add_run(", X = ")
    p33._element.append(make_omml_matrix([['x'], ['y'], ['z']], bracket='['))
    r3 = p33.add_run(", B = ")
    p33._element.append(make_omml_matrix([['6'], ['11'], ['0']], bracket='['))
    r4 = p33.add_run("\n|A| = 1(7) − 1(−3) + 1(−1) = 9 ≠ 0  [1 mark]\nadj A = ")
    p33._element.append(make_omml_matrix([['7', '-3', '2'], ['3', '0', '-3'], ['-1', '3', '1']], bracket='['))
    r5 = p33.add_run("  [2 marks]\nX = A⁻¹B = (1/9) adj(A) B = ")
    p33._element.append(make_omml_matrix([['1'], ['2'], ['3']], bracket='['))
    r6 = p33.add_run("  [2 marks]\nFinal Answer: x = 1, y = 2, z = 3")

    add_p(doc, "Q34. y = (sin⁻¹x)², prove (1 − x²) y₂ − x y₁ − 2 = 0:\ny₁ = 2 sin⁻¹x / √(1 − x²) ⇒ √(1 − x²) y₁ = 2 sin⁻¹x [2 marks]\nDifferentiating: √(1 − x²) y₂ − (x y₁)/√(1 − x²) = 2/√(1 − x²) [2 marks]\nMultiplying by √(1 − x²): (1 − x²) y₂ − x y₁ − 2 = 0. Hence Proved. [1 mark]", space_after=3)
    add_p(doc, "Q35. Evaluate ∫ (x² + 1)/(x² − 5x + 6) dx:\n(x² + 1)/(x² − 5x + 6) = 1 + (5x − 5)/((x − 2)(x − 3)) [1 mark]\nPartial fractions: (5x − 5)/((x − 2)(x − 3)) = −5/(x − 2) + 10/(x − 3) [2 marks]\n∫ [1 − 5/(x − 2) + 10/(x − 3)] dx = x − 5 ln|x − 2| + 10 ln|x − 3| + C [2 marks]\nFinal Answer: x − 5 ln|x − 2| + 10 ln|x − 3| + C", space_after=3)

    add_p(doc, "SECTION E (3 × 4 = 12 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q36. Case Study 1 — Relations and Functions:\n(i) Reflexive: |a − a| = 0 (even) ⇒ (a, a) ∈ R. Symmetric: |a − b| even ⇒ |b − a| even ⇒ (b, a) ∈ R. Transitive: a − b even and b − c even ⇒ a − c even ⇒ |a − c| even ⇒ (a, c) ∈ R. Hence R is equivalence. [2 marks]\n(ii) Equivalence class of 2: [2] = {2, 4} [1 mark]\n(iii) |1 − 5| = 4 (even) ⇒ Yes, (1, 5) ∈ R. [1 mark]", space_after=3)
    add_p(doc, "Q37. Case Study 2 — Application of Derivatives:\n(i) MC(x) = dC/dx = 0.015x² − 0.04x + 30 [2 marks]\n(ii) MC(3) = 0.015(9) − 0.04(3) + 30 = ₹ 30.015 [1 mark]\n(iii) Since MC(3) = 30.015 > 0, cost is increasing at x = 3. [1 mark]", space_after=3)
    add_p(doc, "Q38. Case Study 3 — Integrals:\n(i) s(t) = ∫ (3t² − 12t + 9) dt = t³ − 6t² + 9t + C. With s(0) = 0 ⇒ s(t) = t³ − 6t² + 9t [2 marks]\n(ii) s(2) = 8 − 24 + 18 = 2 m [1 mark]\n(iii) v(t) = 0 ⇒ 3(t − 1)(t − 3) = 0 ⇒ t = 1 s, 3 s [1 mark]", space_after=4)

    add_p(doc, "— End of Marking Scheme —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_a = os.path.join(OUT_DIR, "Set_A_Marking_Scheme.docx")
    doc.save(out_a)
    print(f"Generated {out_a}")


def build_set_b_ms():
    doc = create_base_doc()
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "CLASS XII — MATHEMATICS (SET B)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_p(doc, "MARKING SCHEME & STEP-BY-STEP SOLUTIONS", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Time Allowed: 3 Hours\t\t\t\t\t\t\t\tMaximum Marks: 80", bold=True, size=12, space_after=3)
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "SECTION A (20 Marks — 1 Mark Each)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    mcqs = [
        ("Q1. (c) An equivalence relation\t\t[1 mark]\n    Reason: Reflexive: (1,1),(2,2),(3,3) ∈ R. Symmetric: (1,3) & (3,1) ∈ R. Transitive holds."),
        ("Q2. (b) 6\t\t[1 mark]\n    Reason: Number of one-one functions = 3! = 6."),
        ("Q3. (b) 3π/4\t\t[1 mark]\n    Reason: cos(3π/4) = -1/√2 and 3π/4 ∈ [0, π]."),
        ("Q4. (b) (−π/2, π/2)\t\t[1 mark]\n    Reason: Principal value branch of tan⁻¹x is (-π/2, π/2)."),
        ("Q5. (a) Aᵀ + Bᵀ\t\t[1 mark]\n    Reason: Transpose of sum property."),
        ("Q6. (a) 1\t\t[1 mark]\n    Reason: a₁₂ = -a₂₁ ⇒ a = -2; a₁₃ = -a₃₁ ⇒ -3 = -b ⇒ b = 3. a + b = 1."),
        ("Q7. (b) aᵢⱼ = 0 for i ≠ j\t\t[1 mark]\n    Reason: Definition of diagonal matrix."),
        ("Q8. (b) −54\t\t[1 mark]\n    Reason: |3A| = 3³ |A| = 27 × (-2) = -54."),
        ("Q9. (b) 15/2 sq units\t\t[1 mark]\n    Reason: Area = ½|1(0-3) + 6(3-0) + 4(0-0)| = 15/2."),
        ("Q10. (b) 3cos(3x)\t\t[1 mark]\n    Reason: Chain rule: d/dx(sin 3x) = 3 cos 3x."),
        ("Q11. (c) Both continuous and differentiable everywhere on ℝ\t\t[1 mark]\n    Reason: Polynomial functions are continuous and differentiable on ℝ."),
        ("Q12. (a) 3/4\t\t[1 mark]\n    Reason: At x = 2: k(4) = 3 ⇒ k = 3/4."),
        ("Q13. (a) 4πr²\t\t[1 mark]\n    Reason: dV/dr = 4πr²."),
        ("Q14. (b) (1, 3)\t\t[1 mark]\n    Reason: f'(x) = 3(x - 1)(x - 3) < 0 for x ∈ (1, 3)."),
        ("Q15. (b) 24π cm²/cm\t\t[1 mark]\n    Reason: dS/dr = 8πr = 24π at r = 3."),
        ("Q16. (b) x = −1\t\t[1 mark]\n    Reason: f'(x) = 6(x - 2)(x + 1) = 0. f''(-1) = -18 < 0 ⇒ local max at x = -1."),
        ("Q17. (a) sin x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q18. (a) sin⁻¹x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q19. (a) Both A and R are true, and R is the correct explanation of A\t\t[1 mark]\n    Reason: LHL ≠ RHL at integers, limit does not exist."),
        ("Q20. (a) Both A and R are true, and R is the correct explanation of A\t\t[1 mark]\n    Reason: f'(x) = 1/x > 0 for all x > 0 implies strictly increasing.")
    ]
    for m in mcqs:
        add_p(doc, m, space_after=2)

    add_p(doc, "SECTION B (5 × 2 = 10 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q21. Equivalence check for R = {(1,1),(2,2),(3,3),(1,3),(3,1)}:\n- Reflexive: (1,1),(2,2),(3,3) ∈ R ⇒ Reflexive. [½ mark]\n- Symmetric: (1,3) ∈ R and (3,1) ∈ R ⇒ Symmetric. [½ mark]\n- Transitive: All chains verified ⇒ Transitive. [½ mark]\nConclusion: R is an equivalence relation. [½ mark]", space_after=3)
    
    # Q22 MS with OMML 2x2
    p22 = doc.add_paragraph()
    p22.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p22.paragraph_format.line_spacing = None
    p22.paragraph_format.space_after = Pt(2)
    r = p22.add_run("Q22. A + 2B:\n2B = ")
    p22._element.append(make_omml_matrix([['2', '0'], ['-2', '4']], bracket='['))
    r2 = p22.add_run("  [½ mark]\nA + 2B = ")
    p22._element.append(make_omml_matrix([['5', '1'], ['-3', '6']], bracket='['))
    r3 = p22.add_run("  [1.5 marks]\nFinal Answer: ")
    p22._element.append(make_omml_matrix([['5', '1'], ['-3', '6']], bracket='['))

    add_p(doc, "Q23. Differentiate y = cos(x²):\ndy/dx = −2x sin(x²) [2 marks]\nFinal Answer: −2x sin(x²)", space_after=3)
    add_p(doc, "Q24. Rate of increase of circle area:\nA = πr² ⇒ dA/dt = 2πr (dr/dt) [1 mark]\nWith dr/dt = 3 cm/s and r = 10 cm:\ndA/dt = 2π(10)(3) = 60π cm²/s [1 mark]\nFinal Answer: 60π cm²/s", space_after=3)
    add_p(doc, "Q25. Evaluate ∫ e^(sin x) cos x dx:\nLet u = sin x ⇒ du = cos x dx [1 mark]\n∫ e^(sin x) cos x dx = e^(sin x) + C [1 mark]\nFinal Answer: e^(sin x) + C", space_after=3)

    add_p(doc, "SECTION C (6 × 3 = 18 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q26. Principal values:\n(i) cos⁻¹(cos(5π/4)) = cos⁻¹(cos(2π − 3π/4)) = 3π/4 [1 mark]\n(ii) sin⁻¹(sin(−π/3)) = −π/3 [1 mark]\n(iii) tan⁻¹(tan(5π/6)) = tan⁻¹(−tan(π/6)) = −π/6 [1 mark]", space_after=3)
    add_p(doc, "Q27. Domains & Principal value:\n(i) sin⁻¹(3x − 1): −1 ≤ 3x − 1 ≤ 1 ⇒ Domain: [0, 2/3] [1 mark]\n(ii) cos⁻¹((x − 1)/2): −1 ≤ (x − 1)/2 ≤ 1 ⇒ Domain: [−1, 3] [1 mark]\n(iii) tan⁻¹(−1) = −π/4 [1 mark]", space_after=3)
    
    # Q28 MS with OMML 3x3 Determinant
    p28 = doc.add_paragraph()
    p28.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p28.paragraph_format.line_spacing = None
    p28.paragraph_format.space_after = Pt(2)
    r = p28.add_run("Q28. Determinant expansion along R₁ for |A| = ")
    p28._element.append(make_omml_matrix([['2', '-1', '3'], ['1', '2', '-1'], ['0', '3', '2']], bracket='|'))
    r2 = p28.add_run(":\n|A| = 2(4 + 3) − (−1)(2 − 0) + 3(3 − 0) [2 marks]\n    = 14 + 2 + 9 = 25 [1 mark]\nFinal Answer: 25")

    add_p(doc, "Q29. y = x^(sin x) (x > 0):\nln y = sin x ln x [½ mark]\n(1/y) dy/dx = cos x ln x + (sin x)/x [1.5 marks]\ndy/dx = x^(sin x) [cos x ln x + (sin x)/x] [1 mark]\nFinal Answer: x^(sin x) [cos x ln x + (sin x)/x]", space_after=3)
    add_p(doc, "Q30. Intervals of increase/decrease for f(x) = x³ − 6x² + 9x + 15:\nf'(x) = 3(x − 1)(x − 3) [1 mark]\n- In (−∞, 1): f'(x) > 0 ⇒ Strictly Increasing [1 mark]\n- In (1, 3): f'(x) < 0 ⇒ Strictly Decreasing [½ mark]\n- In (3, ∞): f'(x) > 0 ⇒ Strictly Increasing [½ mark]\nFinal Answer: (i) Strictly increasing on (−∞, 1) ∪ (3, ∞); (ii) Strictly decreasing on (1, 3)", space_after=3)
    add_p(doc, "Q31. Evaluate ∫ x cos x dx using integration by parts:\nu = x ⇒ u' = 1;  v' = cos x ⇒ v = sin x [1 mark]\n∫ x cos x dx = x sin x − ∫ sin x dx = x sin x + cos x + C [2 marks]\nFinal Answer: x sin x + cos x + C", space_after=3)

    add_p(doc, "SECTION D (4 × 5 = 20 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    # Q32 MS with OMML 3x3 Matrices
    p32 = doc.add_paragraph()
    p32.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p32.paragraph_format.line_spacing = None
    p32.paragraph_format.space_after = Pt(2)
    r = p32.add_run("Q32. Express A = ")
    p32._element.append(make_omml_matrix([['3', '-2', '5'], ['4', '1', '-3'], ['-1', '2', '1']], bracket='['))
    r2 = p32.add_run(" as P + Q:\nAᵀ = ")
    p32._element.append(make_omml_matrix([['3', '4', '-1'], ['-2', '1', '2'], ['5', '-3', '1']], bracket='['))
    r3 = p32.add_run("  [1 mark]\nP = ½(A + Aᵀ) = ")
    p32._element.append(make_omml_matrix([['3', '1', '2'], ['1', '1', '-1/2'], ['2', '-1/2', '1']], bracket='['))
    r4 = p32.add_run(" (Symmetric)  [1.5 marks]\nQ = ½(A − Aᵀ) = ")
    p32._element.append(make_omml_matrix([['0', '-3', '3'], ['3', '0', '-5/2'], ['-3', '5/2', '0']], bracket='['))
    r5 = p32.add_run(" (Skew-symmetric)  [1.5 marks]\nVerification: Pᵀ = P, Qᵀ = −Q, and P + Q = A.  [1 mark]")

    # Q33 MS with OMML System
    p33 = doc.add_paragraph()
    p33.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p33.paragraph_format.line_spacing = None
    p33.paragraph_format.space_after = Pt(2)
    r = p33.add_run("Q33. Solve system AX = B:\nA = ")
    p33._element.append(make_omml_matrix([['1', '-1', '2'], ['3', '4', '-5'], ['2', '-1', '3']], bracket='['))
    r2 = p33.add_run(", X = ")
    p33._element.append(make_omml_matrix([['x'], ['y'], ['z']], bracket='['))
    r3 = p33.add_run(", B = ")
    p33._element.append(make_omml_matrix([['7'], ['-5'], ['12']], bracket='['))
    r4 = p33.add_run("\n|A| = 1(7) − (−1)(19) + 2(−11) = 4 ≠ 0  [1 mark]\nadj A = ")
    p33._element.append(make_omml_matrix([['7', '1', '-3'], ['-19', '-1', '11'], ['-11', '-1', '7']], bracket='['))
    r5 = p33.add_run("  [2 marks]\nX = A⁻¹B = (1/4) adj(A) B = ")
    p33._element.append(make_omml_matrix([['2'], ['1'], ['3']], bracket='['))
    r6 = p33.add_run("  [2 marks]\nFinal Answer: x = 2, y = 1, z = 3")

    add_p(doc, "Q34. y = (tan⁻¹x)², show (1 + x²)² y₂ + 2x(1 + x²) y₁ = 2:\ny₁ = 2 tan⁻¹x / (1 + x²) ⇒ (1 + x²) y₁ = 2 tan⁻¹x [2 marks]\nDifferentiating: (1 + x²) y₂ + 2x y₁ = 2 / (1 + x²) [2 marks]\nMultiplying by (1 + x²): (1 + x²)² y₂ + 2x(1 + x²) y₁ = 2. Hence Proved. [1 mark]", space_after=3)
    add_p(doc, "Q35. Evaluate ∫ (x² + x + 1)/((x + 1)(x + 2)) dx:\n(x² + x + 1)/(x² + 3x + 2) = 1 + (−2x − 1)/((x + 1)(x + 2)) [1 mark]\nPartial fractions: (−2x − 1)/((x + 1)(x + 2)) = 1/(x + 1) − 3/(x + 2) [2 marks]\n∫ [1 + 1/(x + 1) − 3/(x + 2)] dx = x + ln|x + 1| − 3 ln|x + 2| + C [2 marks]\nFinal Answer: x + ln|x + 1| − 3 ln|x + 2| + C", space_after=3)

    add_p(doc, "SECTION E (3 × 4 = 12 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q36. Case Study 1 — Relations and Functions:\n(i) Reflexive: a divides a ⇒ (a, a) ∈ R. Transitive: a|b and b|c ⇒ a|c ⇒ (a, c) ∈ R. [2 marks]\n(ii) R is not symmetric: (1, 2) ∈ R (1 divides 2) but (2, 1) ∉ R (2 does not divide 1). [1 mark]\n(iii) R is not an equivalence relation because it is not symmetric. [1 mark]", space_after=3)
    add_p(doc, "Q37. Case Study 2 — Application of Derivatives:\n(i) MR(x) = dR/dx = 26x + 26 [2 marks]\n(ii) MR(7) = 26(7) + 26 = ₹ 208 [1 mark]\n(iii) Since MR(7) = 208 > 0, total revenue is increasing at x = 7. [1 mark]", space_after=3)
    add_p(doc, "Q38. Case Study 3 — Integrals:\n(i) s(t) = ∫ (6t² − 4t) dt = 2t³ − 2t² + C. With s(0) = 0 ⇒ s(t) = 2t³ − 2t² [2 marks]\n(ii) s(3) = 2(27) − 2(9) = 36 m [1 mark]\n(iii) v(t) = 0 ⇒ 2t(3t − 2) = 0 ⇒ t = 2/3 s (for t > 0) [1 mark]", space_after=4)

    add_p(doc, "— End of Marking Scheme —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_b = os.path.join(OUT_DIR, "Set_B_Marking_Scheme.docx")
    doc.save(out_b)
    print(f"Generated {out_b}")


if __name__ == '__main__':
    build_set_a_qp()
    build_set_b_qp()
    build_set_a_ms()
    build_set_b_ms()
    print("All Word documents successfully generated with expanded matrix line height!")
