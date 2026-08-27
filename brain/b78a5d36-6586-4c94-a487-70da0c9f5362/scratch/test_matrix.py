"""
Test inserting native Word OMML equation matrices and determinants.
"""

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_omml_matrix(rows, bracket='['):
    """
    rows: list of lists of strings, e.g. [['1', '2'], ['3', '4']]
    bracket: '[' for matrix, '|' for determinant, '(' for paren matrix
    """
    if bracket == '[':
        beg_chr, end_chr = '[', ']'
    elif bracket == '|':
        beg_chr, end_chr = '|', '|'
    elif bracket == '(':
        beg_chr, end_chr = '(', ')'
    else:
        beg_chr, end_chr = '[', ']'
        
    mr_xml_list = []
    for row in rows:
        e_xml_list = []
        for cell in row:
            e_xml_list.append(f'<m:e><m:r><m:t>{cell}</m:t></m:r></m:e>')
        mr_xml_list.append(f'<m:mr>{"".join(e_xml_list)}</m:mr>')
        
    m_body = "".join(mr_xml_list)
    
    xml = f"""
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:d>
        <m:dPr>
          <m:begChr m:val="{beg_chr}"/>
          <m:endChr m:val="{end_chr}"/>
          <m:grow m:val="on"/>
        </m:dPr>
        <m:e>
          <m:m>
            <m:mPr>
              <m:baseJc m:val="center"/>
              <m:plcHide m:val="on"/>
            </m:mPr>
            {m_body}
          </m:m>
        </m:e>
      </m:d>
    </m:oMath>
    """
    return parse_xml(xml)

doc = Document()
p = doc.add_paragraph("Let A = ")
omml = create_omml_matrix([['0', '1', '-2'], ['-1', '0', 'x'], ['2', '-3', '0']], bracket='[')
p._element.append(omml)
p.add_run(" be a skew-symmetric matrix.")

p2 = doc.add_paragraph("Evaluate the determinant: ")
omml_det = create_omml_matrix([['1', '2', '3'], ['0', '1', '4'], ['5', '6', '0']], bracket='|')
p2._element.append(omml_det)

doc.save(r"C:\Users\rajan\.gemini\antigravity\scratch\test_matrix.docx")
print("Saved test_matrix.docx successfully!")
