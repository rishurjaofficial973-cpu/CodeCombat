"""
Generate Word (.docx) files for Class XII Mathematics Half-Yearly Exam 2026-27.
Reads content from text files and produces properly formatted Word documents.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

DIR = r"C:\Users\rajan\.gemini\antigravity\scratch\math_exam_2026"


def setup_document():
    """Create a document with proper formatting: A4, narrow margins, Times New Roman 12."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(14)
    return doc


def txt_to_docx(input_path, output_path):
    """Convert a text file to a formatted Word document."""
    doc = setup_document()

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip('\n')

        is_center = False
        is_bold = False
        font_size = Pt(12)

        stripped = line.strip()

        if stripped.startswith("HALF-YEARLY") or stripped.startswith("CLASS XII"):
            is_center = True
            is_bold = True
            font_size = Pt(14) if "HALF-YEARLY" in stripped else Pt(13)
        elif stripped.startswith("SECTION "):
            is_center = True
            is_bold = True
            font_size = Pt(13)
        elif stripped.startswith("MARKING SCHEME"):
            is_center = True
            is_bold = True
            font_size = Pt(14)
        elif stripped.startswith("Time Allowed") or stripped.startswith("Maximum Marks"):
            is_bold = True
        elif stripped.startswith("General Instructions"):
            is_bold = True
        elif stripped.startswith("-- End of") or stripped.startswith("\u2014 End of"):
            is_center = True
            is_bold = True

        p = doc.add_paragraph()
        if is_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = font_size
        run.bold = is_bold

    doc.save(output_path)
    print(f"  Created: {output_path}")


def main():
    files = [
        ("Set_A_Question_Paper.txt", "Set_A_Question_Paper.docx"),
        ("Set_A_Marking_Scheme.txt", "Set_A_Marking_Scheme.docx"),
        ("Set_B_Question_Paper.txt", "Set_B_Question_Paper.docx"),
        ("Set_B_Marking_Scheme.txt", "Set_B_Marking_Scheme.docx"),
    ]

    print("Converting text files to Word documents...")
    print(f"Directory: {DIR}\\n")

    ok = 0
    for txt, docx in files:
        inp = os.path.join(DIR, txt)
        out = os.path.join(DIR, docx)
        if not os.path.exists(inp):
            print(f"  SKIPPED (not found): {txt}")
            continue
        try:
            txt_to_docx(inp, out)
            ok += 1
        except Exception as e:
            print(f"  ERROR {txt}: {e}")

    print(f"\\nDone! {ok}/{len(files)} files converted.")
    print(f"Output: {DIR}")


if __name__ == "__main__":
    main()
