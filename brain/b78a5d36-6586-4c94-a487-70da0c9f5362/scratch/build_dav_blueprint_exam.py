"""
Complete Examination Paper Generator with:
Exact Official Heading Template:
- Boxed Header:
  DAV INSTITUTIONS, JHARKHAND ZONE-F
  Half Yearly Examination, 2026-27
  Class: XII                               FULL MARKS: 80
  Subject: Mathematics (Set A / Set B)      TIME ALLOWED: 3 Hours
- Followed by '************' asterisks divider
- All questions aligned with D.A.V. Public Schools, Zone-F Blue Print.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

OUT_DIR = r"C:\Users\rajan\.gemini\antigravity\scratch\math_exam_2026"
IMG_DIR = os.path.join(OUT_DIR, "images")


def make_omml_matrix(rows, bracket='['):
    beg_chr = '|' if bracket == '|' else '['
    end_chr = '|' if bracket == '|' else ']'
    mr_xml_list = []
    for row in rows:
        e_xml_list = []
        for cell in row:
            e_xml_list.append(f'<m:e><m:r><m:rPr><m:scr m:val="roman"/><m:sty m:val="p"/></m:rPr><m:t>{cell}</m:t></m:r></m:e>')
        mr_xml_list.append(f'<m:mr>{"".join(e_xml_list)}</m:mr>')
        
    m_body = "".join(mr_xml_list)
    xml = f"""
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:d>
        <m:dPr>
          <m:begChr m:val="{beg_chr}"/>
          <m:endChr m:val="{end_chr}"/>
          <m:grow m:val="on"/>
        </m:dPr>
        <m:e>
          <m:m>
            <m:mPr>
              <m:baseJc m:val="center"/>
              <m:plcHide m:val="on"/>
            </m:mPr>
            {m_body}
          </m:m>
        </m:e>
      </m:d>
    </m:oMath>
    """
    return parse_xml(xml)


def create_base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    return doc


def add_official_header_box(doc, set_name="Set A", is_marking_scheme=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Cm(18.4)
    
    # Border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>\n'
                          f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'  <w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>\n'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)

    # Margins
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>\n'
                      f'  <w:top w:w="160" w:type="dxa"/>\n'
                      f'  <w:bottom w:w="160" w:type="dxa"/>\n'
                      f'  <w:left w:w="240" w:type="dxa"/>\n'
                      f'  <w:right w:w="240" w:type="dxa"/>\n'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

    # Paragraph 1: Main Title
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r1 = p1.add_run("DAV INSTITUTIONS, JHARKHAND ZONE-F")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(14)

    # Paragraph 2: Subtitle
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6 if not is_marking_scheme else 2)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r2 = p2.add_run("Half Yearly Examination, 2026-27")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11.5)

    if is_marking_scheme:
        p_ms = cell.add_paragraph()
        p_ms.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ms.paragraph_format.space_before = Pt(0)
        p_ms.paragraph_format.space_after = Pt(6)
        p_ms.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r_ms = p_ms.add_run("MARKING SCHEME & STEP-BY-STEP SOLUTIONS")
        r_ms.bold = True
        r_ms.font.name = 'Times New Roman'
        r_ms.font.size = Pt(12)

    # Paragraph 3: Class & Full Marks
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(2)
    p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r3_l = p3.add_run("Class: XII")
    r3_l.font.name = 'Times New Roman'
    r3_l.font.size = Pt(11)

    p3.add_run("\t\t\t\t\t\t\t\t\t")
    r3_r = p3.add_run("FULL MARKS: 80")
    r3_r.font.name = 'Times New Roman'
    r3_r.font.size = Pt(11)

    # Paragraph 4: Subject & Time Allowed
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(2)
    p4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r4_l = p4.add_run(f"Subject: Mathematics ({set_name})")
    r4_l.font.name = 'Times New Roman'
    r4_l.font.size = Pt(11)

    p4.add_run("\t\t\t\t\t\t\t")
    r4_r = p4.add_run("TIME ALLOWED: 3 HOURS")
    r4_r.font.name = 'Times New Roman'
    r4_r.font.size = Pt(11)

    # Asterisks Divider
    p_ast = doc.add_paragraph()
    p_ast.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ast.paragraph_format.space_before = Pt(6)
    p_ast.paragraph_format.space_after = Pt(6)
    p_ast.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r_ast = p_ast.add_run("************")
    r_ast.bold = True
    r_ast.font.name = 'Times New Roman'
    r_ast.font.size = Pt(12)


def add_p(doc, text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = None
    if left_indent > 0:
        p.paragraph_format.left_indent = Cm(left_indent)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_matrix_p(doc, text_before="", rows=None, bracket='[', text_after="", space_after=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = None
    if text_before:
        r1 = p.add_run(text_before)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0, 0, 0)
    if rows:
        p._element.append(make_omml_matrix(rows, bracket=bracket))
    if text_after:
        r2 = p.add_run(text_after)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_image(doc, img_name, width_inches=3.6):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = None
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))


def build_set_a_qp():
    doc = create_base_doc()
    
    # Official Box Header
    add_official_header_box(doc, set_name="Set A", is_marking_scheme=False)
    
    # Instructions
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        add_p(doc, f"{i}. {inst}", size=11, space_after=1, left_indent=0.5)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # SECTION A (20 Marks)
    add_p(doc, "SECTION A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 1 to 18 carry 1 mark each. Questions 19 and 20 are Assertion–Reason based, 1 mark each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "1. Let A = {1, 2, 3} and R = {(1,1), (2,2), (3,3), (1,2)} be a relation on A. Then R is:\n(a) An equivalence relation\t\t\t(b) Reflexive and transitive only\n(c) Reflexive and symmetric only\t\t(d) Symmetric and transitive only", space_after=3)
    add_p(doc, "2. The function f : ℝ → ℝ defined by f(x) = x² is:\n(a) One-one and onto\t\t\t(b) One-one but not onto\n(c) Onto but not one-one\t\t\t(d) Neither one-one nor onto", space_after=3)
    add_p(doc, "3. The principal value of sin⁻¹(−1/2) is:\n(a) π/6\t\t(b) −π/6\t\t(c) 5π/6\t\t(d) −5π/6", space_after=3)
    add_p(doc, "4. The domain of cos⁻¹(2x − 1) is:\n(a) [0, 1]\t\t(b) [−1, 1]\t\t(c) [−1, 0]\t\t(d) [0, 2]", space_after=3)
    add_p(doc, "5. The range of the principal value branch of tan⁻¹x is:\n(a) [−π/2, π/2]\t\t(b) (−π/2, π/2)\t\t(c) [0, π]\t\t(d) (0, π)", space_after=3)
    
    add_matrix_p(doc, "6. The value of x for which the matrix A = ",
                 [['0', '1', '-2'], ['-1', '0', 'x'], ['2', '-3', '0']], '[',
                 " is skew-symmetric, is:\n(a) 3\t\t(b) −3\t\t(c) 2\t\t(d) −2", space_after=3)
    add_p(doc, "7. For any square matrix A, the matrix (A + Aᵀ) is always:\n(a) Skew-symmetric\t\t\t(b) Symmetric\n(c) A diagonal matrix\t\t\t(d) A null matrix", space_after=3)
    add_p(doc, "8. If A and B are symmetric matrices of same order, then (AB − BA) is a:\n(a) Symmetric matrix\t\t\t(b) Skew-symmetric matrix\n(c) Null matrix\t\t\t\t(d) Identity matrix", space_after=3)

    add_p(doc, "9. If A is a 3 × 3 matrix such that |A| = 5, then the value of |3A| is:\n(a) 15\t\t(b) 45\t\t(c) 135\t\t(d) 405", space_after=3)
    add_p(doc, "10. If A is a square matrix of order 3 with |A| = 4, then |adj A| equals:\n(a) 4\t\t(b) 8\t\t(c) 16\t\t(d) 64", space_after=3)

    add_p(doc, "11. If y = e²ˣ, then dy/dx equals:\n(a) e²ˣ\t\t(b) 2e²ˣ\t\t(c) 2x e²ˣ\t\t(d) eˣ", space_after=3)
    add_p(doc, "12. The rate of change of the area of a circle with respect to its radius r, at r = 5 cm, is:\n(a) 5π cm²/cm\t(b) 10π cm²/cm\t(c) 25π cm²/cm\t(d) 2π cm²/cm", space_after=3)

    add_p(doc, "13. ∫ sec²x dx equals:\n(a) tan x + C\t(b) −cot x + C\t(c) sec x tan x + C (d) cot x + C", space_after=3)
    add_p(doc, "14. ∫ 1/(1 + x²) dx equals:\n(a) tan⁻¹x + C\t(b) sin⁻¹x + C\t(c) cot⁻¹x + C\t(d) sec⁻¹x + C", space_after=3)
    add_p(doc, "15. ∫ 1/√(1 − x²) dx equals:\n(a) sin⁻¹x + C\t(b) cos⁻¹x + C\t(c) tan⁻¹x + C\t(d) sec⁻¹x + C", space_after=3)
    add_p(doc, "16. ∫ e³ˣ dx equals:\n(a) e³ˣ + C\t(b) 3e³ˣ + C\t(c) (1/3)e³ˣ + C\t(d) eˣ + C", space_after=3)
    add_p(doc, "17. ∫ cos(2x) dx equals:\n(a) sin(2x) + C\t(b) (1/2)sin(2x) + C\t(c) −2sin(2x) + C\t(d) −(1/2)sin(2x) + C", space_after=3)
    add_p(doc, "18. ∫ (1/x) dx (for x > 0) equals:\n(a) 1/x² + C\t(b) −1/x² + C\t(c) ln(x) + C\t(d) x + C", space_after=3)

    add_p(doc, "19. The graph of a function y = f(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_a_q19.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = |x − 2| shown in the graph is continuous at x = 2.\nReason (R): The function f(x) = |x − 2| is differentiable at x = 2.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=3)

    add_p(doc, "20. The graph of a function y = f(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_a_q20.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The point (3, 4) marked on the graph of f(x) = −(x − 3)² + 4 is a point of local maximum.\nReason (R): At a point of local maximum, f′(x) = 0 and f″(x) < 0.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=5)

    # SECTION B (5 × 2 = 10 Marks)
    add_p(doc, "SECTION B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 21 to 25 carry 2 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "21. Check whether the relation R = {(1,1), (2,2), (3,3), (1,2), (2,3)} on the set A = {1, 2, 3} is reflexive, symmetric and transitive. Justify your answer.", space_after=3)
    add_p(doc, "22. Find the principal value of: tan⁻¹(√3) − sec⁻¹(−2).", space_after=3)
    add_p(doc, "23. If f(x) = kx + 1 for x ≤ 5 and f(x) = 3x − 5 for x > 5 is continuous at x = 5, find the value of k.", space_after=3)
    add_p(doc, "24. Differentiate y = sin(x² + 1) with respect to x.", space_after=3)
    add_p(doc, "25. Evaluate: ∫ x e^(x²) dx.", space_after=3)

    # SECTION C (6 × 3 = 18 Marks)
    add_p(doc, "SECTION C", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 26 to 31 carry 3 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "26. Find the value of:\n(i) sin⁻¹(sin(2π/3))\t\t(ii) cos⁻¹(cos(7π/6))\t\t(iii) tan⁻¹(tan(3π/4))", space_after=3)
    
    p27 = doc.add_paragraph()
    p27.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p27.paragraph_format.space_after = Pt(3)
    p27.add_run("27. If A = ")
    p27._element.append(make_omml_matrix([['1', '2'], ['3', '4']], bracket='['))
    p27.add_run(" and B = ")
    p27._element.append(make_omml_matrix([['2', '0'], ['1', '3']], bracket='['))
    p27.add_run(", find (2A − 3B) and verify that (A + B)ᵀ = Aᵀ + Bᵀ.")

    add_matrix_p(doc, "28. Using minors and cofactors, evaluate the determinant:\n    |A| = ",
                 [['1', '2', '3'], ['0', '1', '4'], ['5', '6', '0']], '|', "", space_after=3)

    add_p(doc, "29. If y = xˣ (x > 0), find dy/dx.", space_after=3)
    add_p(doc, "30. Find the intervals in which the function f(x) = 2x³ − 15x² + 36x + 1 is (i) strictly increasing (ii) strictly decreasing.", space_after=3)
    add_p(doc, "31. Evaluate ∫ x sin x dx using integration by parts.", space_after=3)

    # SECTION D (4 × 5 = 20 Marks)
    add_p(doc, "SECTION D", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 32 to 35 carry 5 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    p32 = doc.add_paragraph()
    p32.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p32.paragraph_format.space_after = Pt(3)
    p32.add_run("32. For the system of linear equations:\n    x + y + z = 6\n    y + 3z = 11\n    x − 2y + z = 0\nwrite the coefficient matrix A = ")
    p32._element.append(make_omml_matrix([['1', '1', '1'], ['0', '1', '3'], ['1', '-2', '1']], bracket='['))
    p32.add_run(". Evaluate the determinant |A|, find adj(A), determine A⁻¹ = adj(A)/|A|, and hence solve the system of equations.")

    add_p(doc, "33. If y = (sin⁻¹x)², prove that (1 − x²) d²y/dx² − x dy/dx − 2 = 0.", space_after=3)
    add_p(doc, "34. An open topped box is to be constructed by removing equal squares of side x from each corner of a 24 cm × 9 cm rectangular sheet of tin and folding up the sides. Find the value of x for which the volume of the box is maximum, and determine this maximum volume.", space_after=3)
    add_p(doc, "35. Evaluate: ∫ (x² + 1)/(x² − 5x + 6) dx.", space_after=3)

    # SECTION E (3 × 4 = 12 Marks)
    add_p(doc, "SECTION E", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 36 to 38 are case-based questions of 4 marks each, with sub-parts of 2, 1 and 1 marks.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q36 Case Study 1 (R&F)
    add_p(doc, "36. Case Study 1 — Relations and Functions", bold=True, size=12, space_after=2)
    add_p(doc, "Two friends, Rohan and Simran, are studying equivalence relations on sets. Consider the set A = {1, 2, 3, 4, 5} and the relation R on A defined by R = {(a, b) : |a − b| is even}. The diagram below illustrates the partition of set A into equivalence classes under relation R:", space_after=2)
    add_image(doc, "set_a_q36.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Show that R is an equivalence relation on A.\t\t(2 marks)\n(ii) Write the equivalence class of the element 2.\t\t(1 mark)\n(iii) Is (1, 5) ∈ R? Justify your answer.\t\t(1 mark)", space_after=4)

    # Q37 Case Study 2 (Matrices)
    add_p(doc, "37. Case Study 2 — Matrices", bold=True, size=12, space_after=2)
    add_p(doc, "A school management committee decides to award prizes to students for three core values: Discipline (x), Cleanliness (y), and Regularity (z). Two schools P and Q decide to award their students as per the allocation scheme shown in the table below:", space_after=2)
    add_image(doc, "set_a_q37.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Represent the given problem in the form of a matrix equation AX = B.\t\t(2 marks)\n(ii) If the prize money for Discipline is ₹ 500 and Regularity is ₹ 600, calculate the prize money for Cleanliness.\t\t(1 mark)\n(iii) State one moral/social benefit of recognizing Cleanliness among students.\t\t(1 mark)", space_after=4)

    # Q38 Case Study 3 (Application of Derivatives)
    add_p(doc, "38. Case Study 3 — Application of Derivatives", bold=True, size=12, space_after=2)
    add_p(doc, "An automobile manufacturing company models its commercial performance. The total revenue received (in thousands of ₹) from the sale of x units of cars is given by R(x) = 3x² + 36x + 5, while the total cost of production is given by C(x) = 0.005x³ − 0.02x² + 30x + 5000. The revenue curve is depicted below:", space_after=2)
    add_image(doc, "set_a_q38.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the marginal revenue function, MR(x) = dR/dx, and evaluate it at x = 10 cars.\t\t(2 marks)\n(ii) Find the marginal cost when x = 2 cars.\t\t(1 mark)\n(iii) Write the profit function P(x) in terms of R(x) and C(x).\t\t(1 mark)", space_after=5)

    add_p(doc, "— End of Question Paper —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_docx = os.path.join(OUT_DIR, "Set_A_Question_Paper.docx")
    doc.save(out_docx)
    print(f"Generated {out_docx}")


def build_set_b_qp():
    doc = create_base_doc()
    
    # Official Box Header
    add_official_header_box(doc, set_name="Set B", is_marking_scheme=False)
    
    # Instructions
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        add_p(doc, f"{i}. {inst}", size=11, space_after=1, left_indent=0.5)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    
    # SECTION A
    add_p(doc, "SECTION A", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 1 to 18 carry 1 mark each. Questions 19 and 20 are Assertion–Reason based, 1 mark each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "1. Let R = {(1,1), (2,2), (3,3), (1,3), (3,1)} be a relation on A = {1, 2, 3}. Then R is:\n(a) Reflexive only\t\t\t(b) Symmetric only\n(c) An equivalence relation\t\t(d) Reflexive and symmetric only", space_after=3)
    add_p(doc, "2. The number of all one-one functions from the set {1, 2, 3} to itself is:\n(a) 3\t\t(b) 6\t\t(c) 9\t\t(d) 27", space_after=3)
    add_p(doc, "3. The principal value of cos⁻¹(−1/√2) is:\n(a) π/4\t\t(b) 3π/4\t\t(c) −π/4\t\t(d) 5π/4", space_after=3)
    add_p(doc, "4. The domain of sin⁻¹(3x − 1) is:\n(a) [0, 2/3]\t\t(b) [−1, 1]\t\t(c) [0, 1]\t\t(d) [−2/3, 2/3]", space_after=3)
    add_p(doc, "5. The principal value of tan⁻¹(−1) is:\n(a) π/4\t\t(b) −π/4\t\t(c) 3π/4\t\t(d) −3π/4", space_after=3)
    
    add_matrix_p(doc, "6. If the matrix A = ",
                 [['0', 'a', '-3'], ['2', '0', '-1'], ['b', '1', '0']], '[',
                 " is skew-symmetric, then the value of (a + b) is:\n(a) 1\t\t(b) −1\t\t(c) 5\t\t(d) −5", space_after=3)
    add_p(doc, "7. A matrix A = [aᵢⱼ]₃ₓ₃ is a diagonal matrix if:\n(a) aᵢⱼ = 0 for all i, j\t\t\t(b) aᵢⱼ = 0 for i ≠ j\n(c) aᵢⱼ = 0 for i = j\t\t\t(d) aᵢⱼ ≠ 0 for all i, j", space_after=3)
    add_p(doc, "8. If A and B are square matrices of the same order, then (A + B)ᵀ equals:\n(a) Aᵀ + Bᵀ\t\t(b) Aᵀ − Bᵀ\t\t(c) AᵀBᵀ\t\t(d) BᵀAᵀ", space_after=3)

    add_p(doc, "9. If A is a 3 × 3 matrix with |A| = −2, then |3A| equals:\n(a) −6\t\t(b) −54\t\t(c) 54\t\t(d) −18", space_after=3)
    add_p(doc, "10. The area of the triangle with vertices (1, 0), (6, 0) and (4, 3), found using determinants, is:\n(a) 15 sq units\t\t(b) 15/2 sq units\t\t(c) 7 sq units\t\t(d) 12 sq units", space_after=3)

    add_p(doc, "11. If y = sin(3x), then dy/dx is:\n(a) cos(3x)\t\t(b) 3cos(3x)\t\t(c) −3cos(3x)\t\t(d) 3sin(3x)", space_after=3)
    add_p(doc, "12. The rate of change of the volume V = (4/3)πr³ of a sphere with respect to its radius r is:\n(a) 4πr²\t\t(b) 2πr\t\t(c) (4/3)πr²\t\t(d) 4πr", space_after=3)

    add_p(doc, "13. ∫ cos x dx equals:\n(a) sin x + C\t(b) −sin x + C\t(c) cos x + C\t(d) −cos x + C", space_after=3)
    add_p(doc, "14. ∫ 1/√(a² − x²) dx equals:\n(a) sin⁻¹(x/a) + C (b) (1/a)sin⁻¹(x/a) + C (c) cos⁻¹(x/a) + C (d) tan⁻¹(x/a) + C", space_after=3)
    add_p(doc, "15. ∫ e⁻ˣ dx equals:\n(a) e⁻ˣ + C\t(b) −e⁻ˣ + C\t(c) eˣ + C\t(d) −eˣ + C", space_after=3)
    add_p(doc, "16. ∫ csc²x dx equals:\n(a) cot x + C\t(b) −cot x + C\t(c) tan x + C\t(d) −csc x + C", space_after=3)
    add_p(doc, "17. ∫ sec x tan x dx equals:\n(a) sec x + C\t(b) tan x + C\t(c) −sec x + C\t(d) cot x + C", space_after=3)
    add_p(doc, "18. ∫ x⁴ dx equals:\n(a) 4x³ + C\t(b) x⁵/5 + C\t(c) 5x⁵ + C\t(d) x³/3 + C", space_after=3)

    add_p(doc, "19. The graph of the greatest integer function y = [x] is shown below:", bold=False, space_after=2)
    add_image(doc, "set_b_q19.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = [x] (greatest integer function) is discontinuous at every integer.\nReason (R): For any integer n, lim(x→n) [x] does not exist.\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=3)

    add_p(doc, "20. The graph of the logarithmic function y = ln(x) is shown below:", bold=False, space_after=2)
    add_image(doc, "set_b_q20.png", width_inches=3.5)
    add_p(doc, "Assertion (A): The function f(x) = log x is strictly increasing on (0, ∞).\nReason (R): f′(x) = 1/x > 0 for all x ∈ (0, ∞).\n(a) Both A and R are true, and R is the correct explanation of A.\n(b) Both A and R are true, but R is not the correct explanation of A.\n(c) A is true, but R is false.\n(d) A is false, but R is true.", space_after=5)

    # SECTION B (5 × 2 = 10 Marks)
    add_p(doc, "SECTION B", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 21 to 25 carry 2 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "21. Let A = {1, 2, 3} and R = {(1,1), (2,2), (3,3), (1,3), (3,1)} be a relation on A. Is R an equivalence relation? Justify your answer.", space_after=3)
    add_p(doc, "22. Find the principal value of: cos⁻¹(−1/2) + 2sin⁻¹(1/2).", space_after=3)
    add_p(doc, "23. The value of k for which f(x) = kx² (for x ≤ 2) and f(x) = 3 (for x > 2) is continuous at x = 2, is:", space_after=3)
    add_p(doc, "24. Differentiate y = cos(x²) with respect to x.", space_after=3)
    add_p(doc, "25. Evaluate: ∫ e^(sin x) · cos x dx.", space_after=3)

    # SECTION C (6 × 3 = 18 Marks)
    add_p(doc, "SECTION C", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 26 to 31 carry 3 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    add_p(doc, "26. Find the value of:\n(i) cos⁻¹(cos(5π/4))\t\t(ii) sin⁻¹(sin(−π/3))\t\t(iii) tan⁻¹(tan(5π/6))", space_after=3)
    
    p27 = doc.add_paragraph()
    p27.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p27.paragraph_format.space_after = Pt(3)
    p27.add_run("27. If A = ")
    p27._element.append(make_omml_matrix([['3', '1'], ['-1', '2']], bracket='['))
    p27.add_run(" and B = ")
    p27._element.append(make_omml_matrix([['1', '0'], ['-1', '2']], bracket='['))
    p27.add_run(", find (A + 2B) and verify that (A + B)ᵀ = Aᵀ + Bᵀ.")

    add_matrix_p(doc, "28. Evaluate the determinant:\n    |A| = ",
                 [['2', '-1', '3'], ['1', '2', '-1'], ['0', '3', '2']], '|', "", space_after=3)

    add_p(doc, "29. If y = x^(sin x) (x > 0), find dy/dx.", space_after=3)
    add_p(doc, "30. Find the intervals in which the function f(x) = x³ − 6x² + 9x + 15 is (i) strictly increasing (ii) strictly decreasing.", space_after=3)
    add_p(doc, "31. Evaluate ∫ x cos x dx using integration by parts.", space_after=3)

    # SECTION D (4 × 5 = 20 Marks)
    add_p(doc, "SECTION D", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 32 to 35 carry 5 marks each.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    p32 = doc.add_paragraph()
    p32.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p32.paragraph_format.space_after = Pt(3)
    p32.add_run("32. For the system of linear equations:\n    x − y + 2z = 7\n    3x + 4y − 5z = −5\n    2x − y + 3z = 12\nwrite the coefficient matrix A = ")
    p32._element.append(make_omml_matrix([['1', '-1', '2'], ['3', '4', '-5'], ['2', '-1', '3']], bracket='['))
    p32.add_run(". Evaluate the determinant |A|, find adj(A), determine A⁻¹ = adj(A)/|A|, and hence solve the system of equations.")

    add_p(doc, "33. If y = (tan⁻¹x)², show that (1 + x²)² d²y/dx² + 2x(1 + x²) dy/dx = 2.", space_after=3)
    add_p(doc, "34. A rectangular sheet of tin 45 cm × 24 cm is to be made into a box without top, by cutting off squares of side x from each corner and folding up the flaps. Find the value of x for which the volume of the box is maximum, and determine this maximum volume.", space_after=3)
    add_p(doc, "35. Evaluate: ∫ (x² + x + 1)/((x + 1)(x + 2)) dx.", space_after=3)

    # SECTION E (3 × 4 = 12 Marks)
    add_p(doc, "SECTION E", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, "(Question numbers 36 to 38 are case-based questions of 4 marks each, with sub-parts of 2, 1 and 1 marks.)", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    # Q36 Case Study 1 (R&F)
    add_p(doc, "36. Case Study 1 — Relations and Functions", bold=True, size=12, space_after=2)
    add_p(doc, "A mathematics club defines a relation R on the set A = {1, 2, 3, 4, 6} as R = {(a, b) : a divides b, where a, b ∈ A}. The directed graph below illustrates the divisibility relation among the elements:", space_after=2)
    add_image(doc, "set_b_q36.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Show that R is reflexive and transitive.\t\t(2 marks)\n(ii) Is R symmetric? Justify with an example.\t\t(1 mark)\n(iii) Is R an equivalence relation? Give reason.\t\t(1 mark)", space_after=4)

    # Q37 Case Study 2 (Matrices)
    add_p(doc, "37. Case Study 2 — Matrices", bold=True, size=12, space_after=2)
    add_p(doc, "Two schools A and B decided to award prizes to their students for three exemplary values: Honesty (x), Hard Work (y), and Cooperation (z). The prize distribution is organized as shown in Figure 37:", space_after=2)
    add_image(doc, "set_b_q37.png", width_inches=4.0)
    add_p(doc, "Based on the above information, answer the following:\n(i) Represent the given problem in the form of a matrix equation AX = B.\t\t(2 marks)\n(ii) If the prize money for Honesty is ₹ 500 and Cooperation is ₹ 400, find the prize money for Hard Work.\t\t(1 mark)\n(iii) State one moral benefit of recognizing Hard Work in schools.\t\t(1 mark)", space_after=4)

    # Q38 Case Study 3 (Application of Derivatives)
    add_p(doc, "38. Case Study 3 — Application of Derivatives", bold=True, size=12, space_after=2)
    add_p(doc, "A manufacturing firm's total revenue (in ₹) received from selling x units of a product is given by R(x) = 13x² + 26x + 15. The revenue function curve is shown below:", space_after=2)
    add_image(doc, "set_b_q37.png", width_inches=3.8)
    add_p(doc, "Based on the above information, answer the following:\n(i) Find the marginal revenue function, MR(x) = dR/dx.\t\t(2 marks)\n(ii) Find the marginal revenue when x = 7 units.\t\t(1 mark)\n(iii) State, with reason, whether the revenue is increasing or decreasing at x = 7.\t\t(1 mark)", space_after=5)

    add_p(doc, "— End of Question Paper —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_docx = os.path.join(OUT_DIR, "Set_B_Question_Paper.docx")
    doc.save(out_docx)
    print(f"Generated {out_docx}")


def build_set_a_ms():
    doc = create_base_doc()
    
    # Official Box Header
    add_official_header_box(doc, set_name="Set A", is_marking_scheme=True)
    
    add_p(doc, "SECTION A (20 Marks — 1 Mark Each)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    mcqs = [
        ("Q1. (b) Reflexive and transitive only\t\t[1 mark]\n    Reason: (1,1),(2,2),(3,3) ∈ R (reflexive). (1,2) ∈ R but (2,1) ∉ R (not symmetric). Transitive condition holds."),
        ("Q2. (d) Neither one-one nor onto\t\t[1 mark]\n    Reason: f(1)=f(-1)=1 (not 1-1). Range = [0, ∞) ≠ ℝ (codomain) (not onto)."),
        ("Q3. (b) −π/6\t\t[1 mark]\n    Reason: sin(-π/6) = -1/2 and -π/6 ∈ [-π/2, π/2]."),
        ("Q4. (a) [0, 1]\t\t[1 mark]\n    Reason: -1 ≤ 2x - 1 ≤ 1 ⇒ 0 ≤ 2x ≤ 2 ⇒ 0 ≤ x ≤ 1."),
        ("Q5. (b) (−π/2, π/2)\t\t[1 mark]\n    Reason: The principal value branch of tan⁻¹x is the open interval (-π/2, π/2)."),
        ("Q6. (a) 3\t\t[1 mark]\n    Reason: a₂₃ = -a₃₂ ⇒ x = -(-3) = 3."),
        ("Q7. (b) Symmetric\t\t[1 mark]\n    Reason: (A + Aᵀ)ᵀ = Aᵀ + A = A + Aᵀ."),
        ("Q8. (b) Skew-symmetric matrix\t\t[1 mark]\n    Reason: (AB − BA)ᵀ = (AB)ᵀ − (BA)ᵀ = BᵀAᵀ − AᵀBᵀ = BA − AB = −(AB − BA)."),
        ("Q9. (c) 135\t\t[1 mark]\n    Reason: |3A| = 3³ |A| = 27 × 5 = 135."),
        ("Q10. (c) 16\t\t[1 mark]\n    Reason: |adj A| = |A|ⁿ⁻¹ = 4³⁻¹ = 16."),
        ("Q11. (b) 2e²ˣ\t\t[1 mark]\n    Reason: d/dx(e²ˣ) = 2e²ˣ."),
        ("Q12. (b) 10π cm²/cm\t\t[1 mark]\n    Reason: dA/dr = 2πr = 10π at r = 5."),
        ("Q13. (a) tan x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q14. (a) tan⁻¹x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q15. (a) sin⁻¹x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q16. (c) (1/3)e³ˣ + C\t\t[1 mark]\n    Reason: Standard formula for exponential integration."),
        ("Q17. (b) (1/2)sin(2x) + C\t\t[1 mark]\n    Reason: ∫ cos(ax) dx = (1/a)sin(ax) + C."),
        ("Q18. (c) ln(x) + C\t\t[1 mark]\n    Reason: Standard logarithmic integral."),
        ("Q19. (c) A is true, but R is false\t\t[1 mark]\n    Reason: f(x) = |x - 2| is continuous everywhere, but not differentiable at the corner x = 2."),
        ("Q20. (a) Both A and R are true, and R is the correct explanation of A\t\t[1 mark]\n    Reason: f'(3) = 0 and f''(3) = -2 < 0 confirms local maximum at (3, 4).")
    ]
    for m in mcqs:
        add_p(doc, m, space_after=2)

    add_p(doc, "SECTION B (5 × 2 = 10 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    add_p(doc, "Q21. Check Reflexive, Symmetric, Transitive for R = {(1,1),(2,2),(3,3),(1,2),(2,3)}:\n- Reflexive: (1,1),(2,2),(3,3) ∈ R ⇒ Reflexive. [½ mark]\n- Symmetric: (1,2) ∈ R but (2,1) ∉ R ⇒ Not symmetric. [½ mark]\n- Transitive: (1,2) ∈ R and (2,3) ∈ R but (1,3) ∉ R ⇒ Not transitive. [1 mark]\nFinal Answer: R is reflexive, but neither symmetric nor transitive.", space_after=3)
    
    add_p(doc, "Q22. tan⁻¹(√3) − sec⁻¹(−2):\ntan⁻¹(√3) = π/3 [½ mark]\nsec⁻¹(−2) = π − sec⁻¹(2) = π − π/3 = 2π/3 [1 mark]\nValue = π/3 − 2π/3 = −π/3 [½ mark]\nFinal Answer: −π/3", space_after=3)

    add_p(doc, "Q23. Continuity at x = 5:\nLHL = f(5) = 5k + 1 [½ mark]\nRHL = 3(5) − 5 = 10 [½ mark]\n5k + 1 = 10 ⇒ 5k = 9 ⇒ k = 9/5 [1 mark]\nFinal Answer: k = 9/5", space_after=3)

    add_p(doc, "Q24. Differentiate y = sin(x² + 1):\ndy/dx = cos(x² + 1) · d/dx(x² + 1) [1 mark]\n      = 2x cos(x² + 1) [1 mark]\nFinal Answer: 2x cos(x² + 1)", space_after=3)

    add_p(doc, "Q25. Evaluate ∫ x e^(x²) dx:\nLet u = x² ⇒ du = 2x dx ⇒ x dx = du/2 [1 mark]\n∫ x e^(x²) dx = ½ ∫ eᵘ du = ½ eᵘ + C = ½ e^(x²) + C [1 mark]\nFinal Answer: ½ e^(x²) + C", space_after=3)

    add_p(doc, "SECTION C (6 × 3 = 18 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q26. Principal value evaluations:\n(i) sin⁻¹(sin(2π/3)) = sin⁻¹(sin(π − π/3)) = π/3 [1 mark]\n(ii) cos⁻¹(cos(7π/6)) = cos⁻¹(cos(2π − 5π/6)) = 5π/6 [1 mark]\n(iii) tan⁻¹(tan(3π/4)) = tan⁻¹(−tan(π/4)) = −π/4 [1 mark]", space_after=3)
    
    p27 = doc.add_paragraph()
    p27.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p27.paragraph_format.space_after = Pt(2)
    p27.add_run("Q27. 2A − 3B and Transpose property:\n2A − 3B = ")
    p27._element.append(make_omml_matrix([['-4', '4'], ['3', '-1']], bracket='['))
    p27.add_run("  [1.5 marks]\nA + B = ")
    p27._element.append(make_omml_matrix([['3', '2'], ['4', '7']], bracket='['))
    p27.add_run(" ⇒ (A + B)ᵀ = ")
    p27._element.append(make_omml_matrix([['3', '4'], ['2', '7']], bracket='['))
    p27.add_run(" = Aᵀ + Bᵀ. Hence verified. [1.5 marks]")

    p28 = doc.add_paragraph()
    p28.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p28.paragraph_format.space_after = Pt(2)
    r = p28.add_run("Q28. Determinant expansion along R₁ for |A| = ")
    p28._element.append(make_omml_matrix([['1', '2', '3'], ['0', '1', '4'], ['5', '6', '0']], bracket='|'))
    r2 = p28.add_run(":\n|A| = 1(0 − 24) − 2(0 − 20) + 3(0 − 5) [2 marks]\n    = −24 + 40 − 15 = 1 [1 mark]\nFinal Answer: 1")

    add_p(doc, "Q29. y = xˣ (x > 0):\nln y = x ln x [½ mark]\n(1/y) dy/dx = 1 + ln x [1.5 marks]\ndy/dx = xˣ(1 + ln x) [1 mark]\nFinal Answer: xˣ(1 + ln x)", space_after=3)
    add_p(doc, "Q30. Intervals of increase/decrease for f(x) = 2x³ − 15x² + 36x + 1:\nf'(x) = 6x² − 30x + 36 = 6(x − 2)(x − 3) [1 mark]\n- In (−∞, 2): f'(x) > 0 ⇒ Strictly Increasing [1 mark]\n- In (2, 3): f'(x) < 0 ⇒ Strictly Decreasing [½ mark]\n- In (3, ∞): f'(x) > 0 ⇒ Strictly Increasing [½ mark]\nFinal Answer: (i) Strictly increasing on (−∞, 2) ∪ (3, ∞); (ii) Strictly decreasing on (2, 3)", space_after=3)
    add_p(doc, "Q31. Evaluate ∫ x sin x dx using integration by parts:\nu = x ⇒ u' = 1;  v' = sin x ⇒ v = −cos x [1 mark]\n∫ x sin x dx = −x cos x − ∫ (−cos x) dx [1 mark]\n             = −x cos x + sin x + C [1 mark]\nFinal Answer: −x cos x + sin x + C", space_after=3)

    add_p(doc, "SECTION D (4 × 5 = 20 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    p32 = doc.add_paragraph()
    p32.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p32.paragraph_format.space_after = Pt(2)
    p32.add_run("Q32. Solve system AX = B using Determinants & Adjoint:\nA = ")
    p32._element.append(make_omml_matrix([['1', '1', '1'], ['0', '1', '3'], ['1', '-2', '1']], bracket='['))
    p32.add_run(", X = ")
    p32._element.append(make_omml_matrix([['x'], ['y'], ['z']], bracket='['))
    p32.add_run(", B = ")
    p32._element.append(make_omml_matrix([['6'], ['11'], ['0']], bracket='['))
    p32.add_run("\nStep 1: Determinant |A| = 1(1 - (-6)) - 1(0 - 3) + 1(0 - 1) = 7 + 3 - 1 = 9 ≠ 0  [1 mark]\nStep 2: Cofactors & adj A = ")
    p32._element.append(make_omml_matrix([['7', '-3', '2'], ['3', '0', '-3'], ['-1', '3', '1']], bracket='['))
    p32.add_run("  [1.5 marks]\nStep 3: A⁻¹ = (1/9) adj(A)  [0.5 mark]\nStep 4: X = A⁻¹B = (1/9) adj(A) B = ")
    p32._element.append(make_omml_matrix([['1'], ['2'], ['3']], bracket='['))
    p32.add_run("  [2 marks]\nFinal Answer: x = 1, y = 2, z = 3")

    add_p(doc, "Q33. y = (sin⁻¹x)², prove (1 − x²) y₂ − x y₁ − 2 = 0:\ny₁ = 2 sin⁻¹x / √(1 − x²) ⇒ √(1 − x²) y₁ = 2 sin⁻¹x [2 marks]\nDifferentiating: √(1 − x²) y₂ − (x y₁)/√(1 − x²) = 2/√(1 − x²) [2 marks]\nMultiplying by √(1 − x²): (1 − x²) y₂ − x y₁ − 2 = 0. Hence Proved. [1 mark]", space_after=3)

    add_p(doc, "Q34. Open box maximum volume:\nV(x) = x(24 − 2x)(9 − 2x) = 4x³ − 66x² + 216x [1 mark]\nV'(x) = 12x² − 132x + 216 = 12(x − 2)(x − 9) = 0 [1.5 marks]\nSince x ∈ (0, 4.5), critical value is x = 2 cm [½ mark]\nV''(x) = 24x − 132 ⇒ V''(2) = 48 − 132 = −84 < 0 (Maximum) [1 mark]\nMax Volume = V(2) = 2(20)(5) = 200 cm³ [1 mark]\nFinal Answer: x = 2 cm, Maximum Volume = 200 cm³", space_after=3)

    add_p(doc, "Q35. Evaluate ∫ (x² + 1)/(x² − 5x + 6) dx:\n(x² + 1)/(x² − 5x + 6) = 1 + (5x − 5)/((x − 2)(x − 3)) [1 mark]\nPartial fractions: (5x − 5)/((x − 2)(x − 3)) = −5/(x − 2) + 10/(x − 3) [2 marks]\n∫ [1 − 5/(x − 2) + 10/(x − 3)] dx = x − 5 ln|x − 2| + 10 ln|x − 3| + C [2 marks]\nFinal Answer: x − 5 ln|x − 2| + 10 ln|x − 3| + C", space_after=3)

    add_p(doc, "SECTION E (3 × 4 = 12 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q36. Case Study 1 — Relations and Functions:\n(i) Reflexive: |a − a| = 0 (even) ⇒ (a, a) ∈ R. Symmetric: |a − b| even ⇒ |b − a| even ⇒ (b, a) ∈ R. Transitive: a − b even and b − c even ⇒ a − c even ⇒ |a − c| even ⇒ (a, c) ∈ R. Hence R is equivalence. [2 marks]\n(ii) Equivalence class of 2: [2] = {2, 4} [1 mark]\n(iii) |1 − 5| = 4 (even) ⇒ Yes, (1, 5) ∈ R. [1 mark]", space_after=3)
    
    p37 = doc.add_paragraph()
    p37.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p37.paragraph_format.space_after = Pt(2)
    p37.add_run("Q37. Case Study 2 — Matrices:\n(i) System in matrix form AX = B:  [2 marks]\n")
    p37._element.append(make_omml_matrix([['3', '2', '4'], ['4', '1', '3']], bracket='['))
    p37._element.append(make_omml_matrix([['x'], ['y'], ['z']], bracket='['))
    p37.add_run(" = ")
    p37._element.append(make_omml_matrix([['4300'], ['3800']], bracket='['))
    p37.add_run("\n(ii) For School P: 3(500) + 2y + 4(600) = 4300 ⇒ 1500 + 2y + 2400 = 4300 ⇒ 2y = 400 ⇒ y = ₹ 200 [1 mark]\n(iii) Cleanliness promotes hygiene, disease prevention, and civic discipline. [1 mark]")

    add_p(doc, "Q38. Case Study 3 — Application of Derivatives:\n(i) MR(x) = dR/dx = 6x + 36 [1 mark]. At x = 10: MR(10) = 6(10) + 36 = ₹ 96 (in thousands) [1 mark]\n(ii) MC(x) = dC/dx = 0.015x² − 0.04x + 30. At x = 2: MC(2) = 0.015(4) − 0.04(2) + 30 = ₹ 29.98 [1 mark]\n(iii) Profit Function P(x) = R(x) − C(x) = (3x² + 36x + 5) − (0.005x³ − 0.02x² + 30x + 5000) [1 mark]", space_after=4)

    add_p(doc, "— End of Marking Scheme —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_a = os.path.join(OUT_DIR, "Set_A_Marking_Scheme.docx")
    doc.save(out_a)
    print(f"Generated {out_a}")


def build_set_b_ms():
    doc = create_base_doc()
    
    # Official Box Header
    add_official_header_box(doc, set_name="Set B", is_marking_scheme=True)
    
    add_p(doc, "SECTION A (20 Marks — 1 Mark Each)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    mcqs = [
        ("Q1. (c) An equivalence relation\t\t[1 mark]\n    Reason: Reflexive: (1,1),(2,2),(3,3) ∈ R. Symmetric: (1,3) & (3,1) ∈ R. Transitive holds."),
        ("Q2. (b) 6\t\t[1 mark]\n    Reason: Number of one-one functions = 3! = 6."),
        ("Q3. (b) 3π/4\t\t[1 mark]\n    Reason: cos(3π/4) = -1/√2 and 3π/4 ∈ [0, π]."),
        ("Q4. (a) [0, 2/3]\t\t[1 mark]\n    Reason: -1 ≤ 3x - 1 ≤ 1 ⇒ 0 ≤ 3x ≤ 2 ⇒ 0 ≤ x ≤ 2/3."),
        ("Q5. (b) −π/4\t\t[1 mark]\n    Reason: tan(-π/4) = -1 and -π/4 ∈ (-π/2, π/2)."),
        ("Q6. (a) 1\t\t[1 mark]\n    Reason: a₁₂ = -a₂₁ ⇒ a = -2; a₁₃ = -a₃₁ ⇒ -3 = -b ⇒ b = 3. a + b = 1."),
        ("Q7. (b) aᵢⱼ = 0 for i ≠ j\t\t[1 mark]\n    Reason: Definition of diagonal matrix."),
        ("Q8. (a) Aᵀ + Bᵀ\t\t[1 mark]\n    Reason: Transpose of sum property."),
        ("Q9. (b) −54\t\t[1 mark]\n    Reason: |3A| = 3³ |A| = 27 × (-2) = -54."),
        ("Q10. (b) 15/2 sq units\t\t[1 mark]\n    Reason: Area = ½|1(0-3) + 6(3-0) + 4(0-0)| = 15/2."),
        ("Q11. (b) 3cos(3x)\t\t[1 mark]\n    Reason: Chain rule: d/dx(sin 3x) = 3 cos 3x."),
        ("Q12. (a) 4πr²\t\t[1 mark]\n    Reason: dV/dr = 4πr²."),
        ("Q13. (a) sin x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q14. (a) sin⁻¹(x/a) + C\t\t[1 mark]\n    Reason: Standard inverse sine integral."),
        ("Q15. (b) −e⁻ˣ + C\t\t[1 mark]\n    Reason: ∫ e^(-x) dx = -e^(-x) + C."),
        ("Q16. (b) −cot x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q17. (a) sec x + C\t\t[1 mark]\n    Reason: Standard formula."),
        ("Q18. (b) x⁵/5 + C\t\t[1 mark]\n    Reason: Power rule: ∫ x^n dx = x^(n+1)/(n+1)."),
        ("Q19. (a) Both A and R are true, and R is the correct explanation of A\t\t[1 mark]\n    Reason: LHL ≠ RHL at integers, limit does not exist."),
        ("Q20. (a) Both A and R are true, and R is the correct explanation of A\t\t[1 mark]\n    Reason: f'(x) = 1/x > 0 for all x > 0 implies strictly increasing.")
    ]
    for m in mcqs:
        add_p(doc, m, space_after=2)

    add_p(doc, "SECTION B (5 × 2 = 10 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q21. Equivalence check for R = {(1,1),(2,2),(3,3),(1,3),(3,1)}:\n- Reflexive: (1,1),(2,2),(3,3) ∈ R ⇒ Reflexive. [½ mark]\n- Symmetric: (1,3) ∈ R and (3,1) ∈ R ⇒ Symmetric. [½ mark]\n- Transitive: All chains verified ⇒ Transitive. [½ mark]\nConclusion: R is an equivalence relation. [½ mark]", space_after=3)
    add_p(doc, "Q22. cos⁻¹(−1/2) + 2sin⁻¹(1/2):\ncos⁻¹(−1/2) = 2π/3 [1 mark]\n2sin⁻¹(1/2) = 2(π/6) = π/3 [½ mark]\nSum = 2π/3 + π/3 = π [½ mark]\nFinal Answer: π", space_after=3)
    add_p(doc, "Q23. Continuity at x = 2:\nk(2²) = 3 ⇒ 4k = 3 ⇒ k = 3/4 [2 marks]\nFinal Answer: k = 3/4", space_after=3)
    add_p(doc, "Q24. Differentiate y = cos(x²):\ndy/dx = −2x sin(x²) [2 marks]\nFinal Answer: −2x sin(x²)", space_after=3)
    add_p(doc, "Q25. Evaluate ∫ e^(sin x) cos x dx:\nLet u = sin x ⇒ du = cos x dx [1 mark]\n∫ e^(sin x) cos x dx = e^(sin x) + C [1 mark]\nFinal Answer: e^(sin x) + C", space_after=3)

    add_p(doc, "SECTION C (6 × 3 = 18 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q26. Principal values:\n(i) cos⁻¹(cos(5π/4)) = cos⁻¹(cos(2π − 3π/4)) = 3π/4 [1 mark]\n(ii) sin⁻¹(sin(−π/3)) = −π/3 [1 mark]\n(iii) tan⁻¹(tan(5π/6)) = tan⁻¹(−tan(π/6)) = −π/6 [1 mark]", space_after=3)
    
    p27 = doc.add_paragraph()
    p27.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p27.paragraph_format.space_after = Pt(2)
    p27.add_run("Q27. A + 2B and Transpose:\nA + 2B = ")
    p27._element.append(make_omml_matrix([['5', '1'], ['-3', '6']], bracket='['))
    p27.add_run("  [1.5 marks]\nA + B = ")
    p27._element.append(make_omml_matrix([['4', '1'], ['-2', '4']], bracket='['))
    p27.add_run(" ⇒ (A + B)ᵀ = ")
    p27._element.append(make_omml_matrix([['4', '-2'], ['1', '4']], bracket='['))
    p27.add_run(" = Aᵀ + Bᵀ. Hence verified. [1.5 marks]")

    p28 = doc.add_paragraph()
    p28.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p28.paragraph_format.space_after = Pt(2)
    r = p28.add_run("Q28. Determinant expansion along R₁ for |A| = ")
    p28._element.append(make_omml_matrix([['2', '-1', '3'], ['1', '2', '-1'], ['0', '3', '2']], bracket='|'))
    r2 = p28.add_run(":\n|A| = 2(4 + 3) − (−1)(2 − 0) + 3(3 − 0) [2 marks]\n    = 14 + 2 + 9 = 25 [1 mark]\nFinal Answer: 25")

    add_p(doc, "Q29. y = x^(sin x) (x > 0):\nln y = sin x ln x [½ mark]\n(1/y) dy/dx = cos x ln x + (sin x)/x [1.5 marks]\ndy/dx = x^(sin x) [cos x ln x + (sin x)/x] [1 mark]\nFinal Answer: x^(sin x) [cos x ln x + (sin x)/x]", space_after=3)
    add_p(doc, "Q30. Intervals of increase/decrease for f(x) = x³ − 6x² + 9x + 15:\nf'(x) = 3(x − 1)(x − 3) [1 mark]\n- In (−∞, 1): f'(x) > 0 ⇒ Strictly Increasing [1 mark]\n- In (1, 3): f'(x) < 0 ⇒ Strictly Decreasing [½ mark]\n- In (3, ∞): f'(x) > 0 ⇒ Strictly Increasing [½ mark]\nFinal Answer: (i) Strictly increasing on (−∞, 1) ∪ (3, ∞); (ii) Strictly decreasing on (1, 3)", space_after=3)
    add_p(doc, "Q31. Evaluate ∫ x cos x dx using integration by parts:\nu = x ⇒ u' = 1;  v' = cos x ⇒ v = sin x [1 mark]\n∫ x cos x dx = x sin x − ∫ sin x dx = x sin x + cos x + C [2 marks]\nFinal Answer: x sin x + cos x + C", space_after=3)

    add_p(doc, "SECTION D (4 × 5 = 20 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    
    p32 = doc.add_paragraph()
    p32.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p32.paragraph_format.space_after = Pt(2)
    p32.add_run("Q32. Solve system AX = B using Determinants & Adjoint:\nA = ")
    p32._element.append(make_omml_matrix([['1', '-1', '2'], ['3', '4', '-5'], ['2', '-1', '3']], bracket='['))
    p32.add_run(", X = ")
    p32._element.append(make_omml_matrix([['x'], ['y'], ['z']], bracket='['))
    p32.add_run(", B = ")
    p32._element.append(make_omml_matrix([['7'], ['-5'], ['12']], bracket='['))
    p32.add_run("\nStep 1: Determinant |A| = 1(7) − (−1)(19) + 2(−11) = 4 ≠ 0  [1 mark]\nStep 2: Cofactors & adj A = ")
    p32._element.append(make_omml_matrix([['7', '1', '-3'], ['-19', '-1', '11'], ['-11', '-1', '7']], bracket='['))
    p32.add_run("  [1.5 marks]\nStep 3: A⁻¹ = (1/4) adj(A)  [0.5 mark]\nStep 4: X = A⁻¹B = (1/4) adj(A) B = ")
    p32._element.append(make_omml_matrix([['2'], ['1'], ['3']], bracket='['))
    p32.add_run("  [2 marks]\nFinal Answer: x = 2, y = 1, z = 3")

    add_p(doc, "Q33. y = (tan⁻¹x)², show (1 + x²)² y₂ + 2x(1 + x²) y₁ = 2:\ny₁ = 2 tan⁻¹x / (1 + x²) ⇒ (1 + x²) y₁ = 2 tan⁻¹x [2 marks]\nDifferentiating: (1 + x²) y₂ + 2x y₁ = 2 / (1 + x²) [2 marks]\nMultiplying by (1 + x²): (1 + x²)² y₂ + 2x(1 + x²) y₁ = 2. Hence Proved. [1 mark]", space_after=3)

    add_p(doc, "Q34. Open box maximum volume:\nV(x) = x(45 − 2x)(24 − 2x) = 4x³ − 138x² + 1080x [1 mark]\nV'(x) = 12x² − 276x + 1080 = 12(x − 5)(x − 18) = 0 [1.5 marks]\nSince x ∈ (0, 12), critical value is x = 5 cm [½ mark]\nV''(x) = 24x − 276 ⇒ V''(5) = 120 − 276 = −156 < 0 (Maximum) [1 mark]\nMax Volume = V(5) = 5(35)(14) = 2450 cm³ [1 mark]\nFinal Answer: x = 5 cm, Maximum Volume = 2450 cm³", space_after=3)

    add_p(doc, "Q35. Evaluate ∫ (x² + x + 1)/((x + 1)(x + 2)) dx:\n(x² + x + 1)/(x² + 3x + 2) = 1 + (−2x − 1)/((x + 1)(x + 2)) [1 mark]\nPartial fractions: (−2x − 1)/((x + 1)(x + 2)) = 1/(x + 1) − 3/(x + 2) [2 marks]\n∫ [1 + 1/(x + 1) − 3/(x + 2)] dx = x + ln|x + 1| − 3 ln|x + 2| + C [2 marks]\nFinal Answer: x + ln|x + 1| − 3 ln|x + 2| + C", space_after=3)

    add_p(doc, "SECTION E (3 × 4 = 12 Marks)", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, "Q36. Case Study 1 — Relations and Functions:\n(i) Reflexive: a divides a ⇒ (a, a) ∈ R. Transitive: a|b and b|c ⇒ a|c ⇒ (a, c) ∈ R. [2 marks]\n(ii) R is not symmetric: (1, 2) ∈ R (1 divides 2) but (2, 1) ∉ R (2 does not divide 1). [1 mark]\n(iii) R is not an equivalence relation because it is not symmetric. [1 mark]", space_after=3)
    
    p37 = doc.add_paragraph()
    p37.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p37.paragraph_format.space_after = Pt(2)
    p37.add_run("Q37. Case Study 2 — Matrices:\n(i) Matrix formulation AX = B:  [2 marks]\n")
    p37._element.append(make_omml_matrix([['4', '3', '2'], ['5', '2', '4']], bracket='['))
    p37._element.append(make_omml_matrix([['x'], ['y'], ['z']], bracket='['))
    p37.add_run(" = ")
    p37._element.append(make_omml_matrix([['3700'], ['4600']], bracket='['))
    p37.add_run("\n(ii) For School A: 4(500) + 3y + 2(400) = 3700 ⇒ 2000 + 3y + 800 = 3700 ⇒ 3y = 900 ⇒ y = ₹ 300 [1 mark]\n(iii) Encouraging Hard Work instills perseverance, resilience, and a growth mindset. [1 mark]")

    add_p(doc, "Q38. Case Study 3 — Application of Derivatives:\n(i) MR(x) = dR/dx = 26x + 26 [2 marks]\n(ii) MR(7) = 26(7) + 26 = 182 + 26 = ₹ 208 [1 mark]\n(iii) Since MR(7) = 208 > 0, total revenue is increasing at x = 7. [1 mark]", space_after=4)

    add_p(doc, "— End of Marking Scheme —", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    out_b = os.path.join(OUT_DIR, "Set_B_Marking_Scheme.docx")
    doc.save(out_b)
    print(f"Generated {out_b}")


if __name__ == '__main__':
    build_set_a_qp()
    build_set_b_qp()
    build_set_a_ms()
    build_set_b_ms()
    print("All documents regenerated with exact DAV Institutions, Jharkhand Zone-F official header box!")
