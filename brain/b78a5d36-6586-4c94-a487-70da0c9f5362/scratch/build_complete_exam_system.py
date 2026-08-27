"""
Build complete Class XII Mathematics Half-Yearly Exam 2026-27 (Set A & Set B).
Includes:
- 100% NCERT Rationalized Syllabus (Zero deleted topics: no tangents/normals, no inverse trig formulas, etc.)
- High-contrast, centrally-aligned embedded diagrams for Case Studies & Assertion-Reason questions
- Proper Word document formatting: Times New Roman 12pt, A4 Portrait, Narrow Margins, Single Spacing, No Boxes
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

BASE_DIR = r"C:\Users\rajan\.gemini\antigravity\scratch\math_exam_2026"
IMG_DIR = os.path.join(BASE_DIR, "images")


def create_base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(14)
    return doc


def add_p(doc, text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_image(doc, img_name, width_inches=4.2):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))


def add_exam_header(doc, title):
    add_p(doc, "HALF-YEARLY EXAMINATION 2026–27", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, title, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("Time Allowed: 3 Hours")
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    
    p.add_run("\t\t\t\t\t\t\t\t")
    r2 = p.add_run("Maximum Marks: 80")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    
    # Separator
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


def add_instructions(doc):
    add_p(doc, "General Instructions:", bold=True, size=12, space_after=2)
    insts = [
        "This question paper contains five sections — A, B, C, D and E. All questions are compulsory.",
        "Section A consists of 20 questions of 1 mark each (Q1–Q20).",
        "Question numbers 19 and 20 are Assertion–Reason based questions. Two statements are given, marked Assertion (A) and Reason (R). Select the correct option from the four given below each question.",
        "Section B consists of 5 questions of 2 marks each (Q21–Q25).",
        "Section C consists of 6 questions of 3 marks each (Q26–Q31).",
        "Section D consists of 4 questions of 5 marks each (Q32–Q35).",
        "Section E consists of 3 case-based/source-based questions of 4 marks each (Q36–Q38), with sub-parts of 2, 1 and 1 marks.",
        "There is no overall choice. All questions are compulsory as set.",
        "Use of calculators is not permitted."
    ]
    for i, inst in enumerate(insts, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{i}. {inst}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
    
    add_p(doc, "—" * 70, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)


def add_section_title(doc, title, sub):
    add_p(doc, title, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    add_p(doc, sub, italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)


print("Script template ready.")
